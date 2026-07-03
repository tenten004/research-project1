# -*- coding: utf-8 -*-
"""
大学院入試 面接資料（研究概要と経過）を Word (.docx) で生成するスクリプト。

実行:
    python 大学院入試/generate_interview_doc.py
出力:
    大学院入試/大学院入試面接_研究概要と経過.docx
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

FONT_JP = "Yu Gothic"
COLOR_MAIN = RGBColor(0x1F, 0x3B, 0x63)   # 濃紺
COLOR_ACCENT = RGBColor(0x2E, 0x74, 0xB5)  # 青
COLOR_TEXT = RGBColor(0x22, 0x22, 0x22)


# --------------------------------------------------------------------------
# フォント・スタイル関連ヘルパー
# --------------------------------------------------------------------------
def _apply_font(run, size=None, bold=None, italic=None, color=None, name=FONT_JP):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _set_base_styles(doc):
    """Normal / 各見出しスタイルに日本語フォントを設定."""
    normal = doc.styles["Normal"]
    normal.font.name = FONT_JP
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = COLOR_TEXT
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_JP)

    for style_name in ("Heading 1", "Heading 2", "Heading 3", "Title", "List Bullet",
                        "List Bullet 2", "List Number", "Quote"):
        try:
            st = doc.styles[style_name]
        except KeyError:
            continue
        st.font.name = FONT_JP
        if st.element.rPr is not None:
            rfonts = st.element.rPr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = st.element.rPr.makeelement(qn("w:rFonts"), {})
                st.element.rPr.append(rfonts)
            rfonts.set(qn("w:eastAsia"), FONT_JP)


# --------------------------------------------------------------------------
# 段落生成ヘルパー
# --------------------------------------------------------------------------
def add_title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    _apply_font(r, size=20, bold=True, color=COLOR_MAIN)
    if subtitle:
        ps = doc.add_paragraph()
        ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rs = ps.add_run(subtitle)
        _apply_font(rs, size=11, bold=False, color=COLOR_ACCENT)


def add_h1(doc, text):
    p = doc.add_heading(level=1)
    r = p.add_run(text)
    _apply_font(r, size=14.5, bold=True, color=COLOR_MAIN)
    return p


def add_h2(doc, text):
    p = doc.add_heading(level=2)
    r = p.add_run(text)
    _apply_font(r, size=12, bold=True, color=COLOR_ACCENT)
    return p


def add_para(doc, text, size=10.5, bold=False, color=COLOR_TEXT, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    _apply_font(r, size=size, bold=bold, color=color)
    return p


def add_bullets(doc, items, size=10.5):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        if isinstance(item, tuple):
            head, body = item
            r1 = p.add_run(head)
            _apply_font(r1, size=size, bold=True, color=COLOR_TEXT)
            r2 = p.add_run(body)
            _apply_font(r2, size=size, bold=False, color=COLOR_TEXT)
        else:
            r = p.add_run(item)
            _apply_font(r, size=size, color=COLOR_TEXT)


def add_qa(doc, question, answer):
    pq = doc.add_paragraph()
    pq.paragraph_format.space_after = Pt(1)
    rq = pq.add_run("Q. " + question)
    _apply_font(rq, size=10.5, bold=True, color=COLOR_ACCENT)
    pa = doc.add_paragraph()
    pa.paragraph_format.space_after = Pt(8)
    pa.paragraph_format.left_indent = Pt(12)
    ra = pa.add_run("A. " + answer)
    _apply_font(ra, size=10.5, color=COLOR_TEXT)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        _apply_font(run, size=9.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            _apply_font(run, size=9.5, color=COLOR_TEXT)
            cells[i].paragraphs[0].alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            )
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_spacer(doc, pt=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pt)


# --------------------------------------------------------------------------
# 本文
# --------------------------------------------------------------------------
def build(doc):
    _set_base_styles(doc)

    add_title(
        doc,
        "大学院入試 面接資料",
        "これまでの研究とその経過 ― MRI 画像による大脳白質病変グレード判別",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("氏名：竹村 典晃")
    _apply_font(r, size=10.5, color=COLOR_TEXT)
    add_spacer(doc, 6)

    # 1. はじめに
    add_h1(doc, "1. はじめに（研究の背景と動機）")
    add_para(
        doc,
        "大脳白質病変は脳に生じる虚血性の変化であり、その最大のリスク因子は高血圧であるとされています。"
        "脳卒中のリスク因子でもあり、早期に進行度（グレード）を把握することは予防医療の観点から重要です。"
        "私が所属する研究室では、これまで脳ドック受診者の健診データを用い、ロジスティック回帰により"
        "大脳白質病変の有無を予測するモデルを構築してきました（先行研究）。",
    )
    add_para(
        doc,
        "私の研究では、この流れを発展させ、健診データではなく MRI 画像そのものから大脳白質病変のグレードを"
        "自動判別する手法の確立を目指しています。学部では畳み込みニューラルネットワーク（CNN）を用いた判別に取り組み、"
        "大学院進学後は Vision Transformer（ViT）を導入して、より汎化性能の高い判別モデルの構築と、"
        "深層学習モデルの適切な評価方法の確立に取り組んでいます。",
    )

    # 2. 研究の全体像
    add_h1(doc, "2. 研究の全体像")
    add_para(
        doc,
        "研究は大きく次の 2 つのフェーズで進めてきました。学部の卒業研究（フェーズ1）で得た知見と課題を踏まえ、"
        "大学院ではモデルと評価設計の両面から研究を深めています（フェーズ2）。",
    )
    add_bullets(
        doc,
        [
            ("フェーズ1（学部・卒業研究）：", "CNN による大脳白質病変のグレード予測。撮像スライスの取り方と判別精度の関係、および健診データとの関連を分析。"),
            ("フェーズ2（大学院・発展研究）：", "ViT の導入と CNN との比較。患者単位分割による厳密な再評価、クラス不均衡への対処、少数グレード検出の改善。"),
        ],
    )

    # 3. 学部での研究
    add_h1(doc, "3. 学部での研究：CNN による大脳白質病変のグレード予測")

    add_h2(doc, "3.1 研究目的")
    add_para(
        doc,
        "MRI 画像を用いた大脳白質病変の診断（グレード判別）手法を確立すること、"
        "および判別結果と健診データの検査項目との関連を明らかにすることを目的としました。",
    )

    add_h2(doc, "3.2 対象データと「アキシャル」の定義")
    add_bullets(
        doc,
        [
            "頭部 MRI は、患者 1 人あたり頭頂部から眼球付近まで、およそ 25 枚の横断面（スライス）画像で構成される。",
            "撮影された各横断面スライスの順番を、便宜的に「アキシャル」という単位として定義した。",
            "大脳白質病変は脳室周囲に多く現れ、頭頂部寄りや眼球側のスライスには病変が写りにくい。",
            "画像には FLAIR 画像・T1 強調画像・T2 強調画像の 3 モダリティが存在し、組み合わせを比較検討した。",
        ],
    )

    add_h2(doc, "3.3 手法")
    add_bullets(
        doc,
        [
            "CNN（畳み込みニューラルネットワーク）でグレード（0〜4 の 5 段階）を判別。",
            "学習・評価に用いるアキシャルの位置と範囲（range）、および入力画像サイズを変化させ、判別性能の変化を系統的に検証。",
            "健診データとの関連分析として、高血圧・糖尿病・脂質異常症・肥満などの項目をカテゴリ変数化し、共起ネットワーク図で可視化。",
        ],
    )

    add_h2(doc, "3.4 主な結果")
    add_para(
        doc,
        "最も高い判別性能を示したのは、入力画像を 80×80 ピクセル、アキシャル 9〜15（range 7）、"
        "FLAIR 画像と T1 強調画像を用いたときで、テストデータでの精度は約 0.92 に達しました。"
        "各グレードの AUC は以下のとおりです。",
    )
    add_table(
        doc,
        ["グレード", "AUC"],
        [
            ["grade0", "0.9814"],
            ["grade1", "0.9800"],
            ["grade2", "0.9905"],
            ["grade3", "0.9977"],
            ["grade4", "0.9998"],
        ],
    )
    add_para(
        doc,
        "また共起ネットワーク図からは、大脳白質病変が進行した患者ほど正常血圧から遠く、II 度・III 度高血圧に近い位置に分布し、"
        "病変の進行度と高血圧の重症度に関連がある傾向が示唆されました。",
    )

    add_h2(doc, "3.5 考察と残された課題")
    add_bullets(
        doc,
        [
            "range（使用アキシャル数）が大きいほど学習画像数が増え精度が上がる一方、病変の写らない頭頂部・眼球側を含めると精度が低下した。",
            "エポック 25 付近から学習データと評価データの精度・損失に差が生じ、過学習の傾向が見られた。",
            "診断自動化の観点では精度になお課題が残る。より幅広い健診項目（降圧剤等の薬剤使用、脳梗塞の既往歴など）との関連検討も今後の課題として残った。",
        ],
    )

    # 4. 大学院での発展
    add_h1(doc, "4. 大学院での研究の発展：ViT の導入と患者単位分割による再検証")

    add_h2(doc, "4.1 問題意識 ― データリーケージの発見")
    add_para(
        doc,
        "研究を発展させる過程で、学部研究の評価設定に重大な問題があることに気付きました。"
        "従来は画像単位でデータを train/test 分割していましたが、大脳白質病変のラベルは「患者単位」で付与されます"
        "（1 人の患者の約 46 枚すべてのスライスが同一グレード。全 1154 患者が単一ラベルであることを確認）。"
        "そのため画像単位分割では、評価データの患者のほぼ全員（検証 289 患者のうち大多数）が学習データにも含まれており、"
        "高い精度は「未知の患者に対する汎化」ではなく「学習済み患者の記憶」を測っていたことが分かりました。",
    )

    add_h2(doc, "4.2 研究の経過（実験ステップ）")
    add_para(doc, "問題の発見後、評価の妥当性を担保しながら、以下の手順で段階的に実験を進めました。")
    add_bullets(
        doc,
        [
            ("① 患者単位分割の導入：", "患者リーケージ 0（学習 865 患者 / 検証 289 患者）で再分割。厳密な評価に切り替えると精度は約 0.92 から約 0.5 へ低下し、真の難易度が明らかになった（多数派ベースライン精度 ≒ 0.48）。"),
            ("② クラス不均衡の把握：", "患者単位のグレード分布は grade0=557・1=444・2=106・3=39・4=8 と極端な不均衡。特に grade4 は全 8 患者しかない。"),
            ("③ 3 クラス統合：", "grade0 / 1 / 2+（2 以上を統合）に再定義したところ、macro-F1 がほぼ倍増（CNN で 0.258 → 0.489）し、少数クラスの予測崩壊を回避できた。"),
            ("④ クラス重み付けの検証：", "class-balanced loss 単独で検証。3 クラス化後は不均衡が緩和済みのため改善は見られず（崩壊もせず）、重み付けは不要と結論。過補正手法の重ねがけが崩壊を招く教訓も得た。"),
            ("⑤ 患者単位集約評価：", "スライス単位の予測を患者単位に集約（mean / max_confidence / attention 等）することで、特に AUC が改善することを確認。"),
            ("⑥ アキシャル 9〜15 への限定：", "病変が写りにくい周辺スライスを除くことで、少数クラス（grade2+）の recall が 0.179 → 0.410（集約法により最大 0.513）へ大幅改善。"),
        ],
    )

    add_h2(doc, "4.3 主要な結果（患者単位・3 クラス・mean 集約）")
    add_table(
        doc,
        ["実験条件", "モデル", "Acc", "macro-F1", "AUC", "grade2+ recall"],
        [
            ["3 クラス 全アキシャル", "CNN (ResNet18)", "0.578", "0.489", "0.721", "0.179"],
            ["3 クラス 全アキシャル", "ViT (vit_base)", "0.519", "0.411", "0.640", "0.103"],
            ["3 クラス アキシャル9-15", "CNN (ResNet18)", "0.564", "0.496", "0.701", "0.410"],
        ],
    )

    add_h2(doc, "4.4 得られた主要な知見と現在の到達点")
    add_bullets(
        doc,
        [
            "評価設計そのものが結論を左右する。患者単位分割の導入により、臨床応用を見据えた「正しい難易度」で議論できるようになった。",
            "5 クラス → 3 クラス統合と周辺スライス除去が、少数クラス検出に有効であることを定量的に示した。",
            "現行の再現条件（80px = 5×5 トークン、RMSprop、data augmentation なし、大型の vit_base）は ViT に構造的に不利であり、全条件で CNN が ViT を上回っている。",
            "そこで現在は、ViT の強みを引き出す公平な設定（224px 入力、AdamW、data augmentation、DeiT/ViT-small 等の軽量バックボーン）で再評価を進めている。",
        ],
    )

    # 5. 今後の展望
    add_h1(doc, "5. 大学院で取り組みたい研究テーマと展望")
    add_bullets(
        doc,
        [
            "ViT に適した学習設定・軽量アーキテクチャの検討により、CNN と公平かつ十分な条件で比較し、Transformer 系モデルの有効性を明らかにする。",
            "Attention 機構を用いて、判別根拠となる脳内部位（病変領域）を可視化し、判別の説明可能性（医療における納得性）を高める。",
            "極端なクラス不均衡・少数症例に頑健な学習手法（適切な集約・損失設計）の確立。",
            "MRI 画像と健診データ（血圧・生化学検査等）を統合したマルチモーダルな予測・リスク評価への拡張。",
            "最終的に、大脳白質病変の診断支援に資する、汎化性能と説明可能性を両立したモデルの構築を目指す。",
        ],
    )

    # 6. 習得した知識・技術
    add_h1(doc, "6. 研究を通じて習得した知識・技術")
    add_bullets(
        doc,
        [
            ("プログラミング／深層学習：", "Python、PyTorch、timm（ViT・DeiT・ResNet の利用）、CNN／Transformer の実装と学習。"),
            ("医療画像処理：", "MRI 画像（FLAIR / T1 / T2）の前処理、リサイズ・正規化、データセット構築。"),
            ("実験設計・評価：", "患者単位分割による厳密な評価、accuracy / macro-F1 / ROC-AUC / 混同行列、クラス不均衡対策、再現性を意識した YAML 設定管理。"),
            ("データ分析・可視化：", "共起ネットワーク図による健診データとの関連分析、結果の可視化。"),
            ("開発運用：", "Git / GitHub によるバージョン管理、実験ログとスナップショットによる結果管理。"),
        ],
    )

    # 7. 想定問答
    add_h1(doc, "7. 面接想定問答")
    add_qa(
        doc,
        "なぜ CNN から ViT へ発展させたのですか？",
        "CNN は局所特徴の抽出に優れる一方、大脳白質病変は脳室周囲の広がりなど大域的な文脈が重要です。"
        "自己注意機構により画像全体の関係を捉えられる ViT が、病変分布の把握に適する可能性があると考えたためです。",
    )
    add_qa(
        doc,
        "患者単位分割にしたら精度が大きく下がりましたが、これは後退ではないですか？",
        "いいえ。従来の高精度は患者リーケージによる過大評価で、臨床で必要な「未知の患者への汎化性能」を測れていませんでした。"
        "正しい評価基準を確立したこと自体が本研究の重要な成果だと考えています。",
    )
    add_qa(
        doc,
        "クラス不均衡にはどう対処しましたか？",
        "まず 3 クラス統合で不均衡を緩和し、macro-F1 をほぼ倍増させました。class-balanced loss も検証しましたが、"
        "3 クラス化後は不要でした。手法を重ねがけすると少数クラスを過剰予測して崩壊するため、段階的な検証を重視しています。",
    )
    add_qa(
        doc,
        "現状 ViT が CNN に及んでいないのに、ViT を使う意義は？",
        "現行の再現条件が ViT に構造的に不利（低解像度・少データ・大型モデル）だからです。"
        "224px 入力・AdamW・data augmentation・軽量バックボーン等、ViT に適した設定での公平な比較を進めており、"
        "その条件下での有効性の検証が今後の中心テーマです。",
    )
    add_qa(
        doc,
        "この研究の臨床的・社会的な意義は何ですか？",
        "大脳白質病変は高血圧や脳卒中と関連し、早期の進行度把握は予防に直結します。"
        "MRI からの自動グレード判別は、読影負担の軽減と診断支援に貢献でき、健診データとの統合でリスク評価の高度化も期待できます。",
    )

    # フッター的な一言
    add_spacer(doc, 8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("― 以上 ―")
    _apply_font(r, size=10.5, color=COLOR_ACCENT)


def main():
    out_dir = Path(__file__).resolve().parent
    out_path = out_dir / "大学院入試面接_研究概要と経過.docx"
    doc = Document()
    # 余白を少し狭めて読みやすく
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Pt(48)
        section.left_margin = section.right_margin = Pt(56)
    build(doc)
    doc.save(str(out_path))
    print(f"saved: {out_path}")


if __name__ == "__main__":
    main()
