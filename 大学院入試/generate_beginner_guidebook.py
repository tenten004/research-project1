# -*- coding: utf-8 -*-
"""MRI/ViT研究を初学者向けに説明するWordガイドブックを生成する。"""

from __future__ import annotations

import tempfile
from pathlib import Path
from typing import Iterable, Sequence

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MRI_ROOT = ROOT / "mri-vit-classification"
OUTPUT_PATH = Path(__file__).resolve().parent / "MRI_ViT研究_初学者ガイドブック_20260717.docx"
DETAIL_FIGURES = MRI_ROOT / "outputs/cv5_all_axial_3class/oof_detailed_analysis"

FONT_JP = "Yu Gothic"
NAVY = RGBColor(0x1F, 0x3B, 0x63)
BLUE = RGBColor(0x2E, 0x74, 0xB5)
GREEN = RGBColor(0x54, 0x8B, 0x54)
TEXT = RGBColor(0x22, 0x22, 0x22)
MUTED = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def apply_font(run, size: float | None = None, bold: bool | None = None,
               color: RGBColor | None = None, italic: bool | None = None) -> None:
    run.font.name = FONT_JP
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("eastAsia", "ascii", "hAnsi"):
        rfonts.set(qn(f"w:{attr}"), FONT_JP)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if italic is not None:
        run.font.italic = italic


def setup_document(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT_JP
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_JP)

    for name, size, color in (
        ("Title", 22, NAVY),
        ("Heading 1", 16, NAVY),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 11, BLUE),
    ):
        style = doc.styles[name]
        style.font.name = FONT_JP
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_JP)

    for name in ("List Bullet", "List Bullet 2", "List Number", "Caption"):
        style = doc.styles[name]
        style.font.name = FONT_JP
        style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_JP)

    section = doc.sections[0]
    section.top_margin = Mm(17)
    section.bottom_margin = Mm(16)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(8)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("MRI×AI研究 初学者ガイドブック")
    apply_font(run, size=8.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("- ")
    apply_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])
    tail = footer.add_run(" -")
    apply_font(tail, size=9, color=MUTED)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(level=level)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(10 if level == 1 else 7)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    apply_font(
        run,
        size=16 if level == 1 else (13 if level == 2 else 11),
        bold=True,
        color=NAVY if level == 1 else BLUE,
    )


def add_para(doc: Document, text: str, bold: bool = False,
             color: RGBColor = TEXT, space_after: float = 5) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.2
    paragraph.paragraph_format.space_after = Pt(space_after)
    run = paragraph.add_run(text)
    apply_font(run, size=10.5, bold=bold, color=color)


def add_bullets(doc: Document, items: Iterable[str | tuple[str, str]], level: int = 1) -> None:
    style = "List Bullet" if level == 1 else "List Bullet 2"
    for item in items:
        paragraph = doc.add_paragraph(style=style)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.15
        if isinstance(item, tuple):
            head, body = item
            head_run = paragraph.add_run(head)
            apply_font(head_run, size=10.5, bold=True)
            body_run = paragraph.add_run(body)
            apply_font(body_run, size=10.5)
        else:
            run = paragraph.add_run(item)
            apply_font(run, size=10.5)


def shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def repeat_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[object]],
              font_size: float = 8.8) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    repeat_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = ""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade_cell(cell, "2E74B5")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(str(header))
        apply_font(run, size=font_size, bold=True, color=WHITE)

    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for column_index, value in enumerate(row):
            cell = cells[column_index]
            cell.text = ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index % 2 == 1:
                shade_cell(cell, "F3F6FA")
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if column_index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(str(value))
            apply_font(run, size=font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_note(doc: Document, title: str, body: str, fill: str = "EAF2F8") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    title_run = paragraph.add_run(title)
    apply_font(title_run, size=10.5, bold=True, color=NAVY)
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.15
    body_run = paragraph.add_run(body)
    apply_font(body_run, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_figure(doc: Document, path: Path, caption: str, width: float = 6.4) -> None:
    if not path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_after = Pt(6)
    run = caption_paragraph.add_run(caption)
    apply_font(run, size=9, color=MUTED, italic=True)


def draw_box(axis, xy: tuple[float, float], size: tuple[float, float], text: str,
             color: str, fontsize: int = 10) -> None:
    x, y = xy
    width, height = size
    box = FancyBboxPatch(
        (x, y), width, height,
        boxstyle="round,pad=0.02,rounding_size=0.025",
        linewidth=1.5,
        edgecolor=color,
        facecolor=color + "22",
    )
    axis.add_patch(box)
    axis.text(x + width / 2, y + height / 2, text, ha="center", va="center", fontsize=fontsize)


def draw_arrow(axis, start: tuple[float, float], end: tuple[float, float]) -> None:
    arrow = FancyArrowPatch(start, end, arrowstyle="-|>", mutation_scale=15,
                            linewidth=1.5, color="#666666")
    axis.add_patch(arrow)


def create_pipeline_figure(path: Path) -> None:
    plt.rcParams["font.family"] = FONT_JP
    fig, axis = plt.subplots(figsize=(12, 3.4))
    axis.set_xlim(0, 12)
    axis.set_ylim(0, 3.4)
    axis.axis("off")
    labels = [
        (0.15, "1,154人の患者\nFL＋T1・全axial", "#2E74B5"),
        (2.55, "患者単位\n5-fold分割", "#70AD47"),
        (4.95, "CNN / ViTで\n各スライスを評価", "#8064A2"),
        (7.35, "信頼度上位5枚を\n平均（top-5）", "#C55A11"),
        (9.75, "患者ごとのOOF予測\n統計比較", "#2E74B5"),
    ]
    for x, text, color in labels:
        draw_box(axis, (x, 1.15), (1.95, 1.1), text, color, fontsize=10)
    for x in (2.15, 4.55, 6.95, 9.35):
        draw_arrow(axis, (x, 1.7), (x + 0.35, 1.7))
    axis.text(6, 2.9, "今回の主解析の流れ", ha="center", va="center",
              fontsize=16, fontweight="bold", color="#1F3B63")
    axis.text(6, 0.45, "患者は評価側に1回だけ登場。手動でaxial範囲を選ばず、全画像を入力する。",
              ha="center", va="center", fontsize=10, color="#555555")
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)


def create_split_figure(path: Path) -> None:
    plt.rcParams["font.family"] = FONT_JP
    fig, axes = plt.subplots(1, 2, figsize=(12, 4.5))
    titles = ["誤り：画像単位split", "正しい：患者単位split"]
    for axis, title in zip(axes, titles):
        axis.set_xlim(0, 10)
        axis.set_ylim(0, 8)
        axis.axis("off")
        axis.set_title(title, fontsize=15, fontweight="bold",
                       color="#C00000" if "誤り" in title else "#548B54")

    left = axes[0]
    for index, y in enumerate((5.9, 4.3, 2.7)):
        color = ("#2E74B5", "#C55A11", "#8064A2")[index]
        draw_box(left, (0.4, y), (2.1, 0.9), f"患者{chr(65 + index)}\n複数スライス", color)
        draw_arrow(left, (2.6, y + 0.45), (3.4, y + 0.45))
        draw_box(left, (3.55, y), (2.2, 0.9), "画像をランダム分割", "#777777", 9)
        draw_arrow(left, (5.85, y + 0.45), (6.45, y + 0.45))
        draw_box(left, (6.6, y + 0.45), (2.6, 0.65), "trainに同じ患者", color, 9)
        draw_box(left, (6.6, y - 0.35), (2.6, 0.65), "validationにも同じ患者", color, 9)
    left.text(5, 1.2, "モデルが患者固有の特徴を覚えられる\n→ 未見患者への性能ではない",
              ha="center", va="center", fontsize=11, color="#C00000", fontweight="bold")

    right = axes[1]
    draw_box(right, (0.5, 5.3), (2.3, 1.0), "患者A・B・C\n全スライス", "#2E74B5")
    draw_box(right, (0.5, 2.8), (2.3, 1.0), "患者D\n全スライス", "#C55A11")
    draw_arrow(right, (2.9, 5.8), (4.0, 5.8))
    draw_arrow(right, (2.9, 3.3), (4.0, 3.3))
    draw_box(right, (4.15, 5.1), (2.2, 1.4), "train\n患者A・B・Cだけ", "#70AD47")
    draw_box(right, (4.15, 2.6), (2.2, 1.4), "validation\n患者Dだけ", "#8064A2")
    draw_box(right, (7.1, 3.8), (2.4, 1.4), "患者の重複 0\n未見患者を評価", "#548B54")
    draw_arrow(right, (6.45, 5.1), (7.0, 4.8))
    draw_arrow(right, (6.45, 3.3), (7.0, 4.1))
    right.text(5, 1.2, "試験問題に、練習で見た患者を混ぜない",
               ha="center", va="center", fontsize=11, color="#548B54", fontweight="bold")
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)


def create_cv_figure(path: Path) -> None:
    plt.rcParams["font.family"] = FONT_JP
    fig, axis = plt.subplots(figsize=(11, 4.8))
    axis.set_xlim(0, 11)
    axis.set_ylim(0, 6.5)
    axis.axis("off")
    axis.text(5.5, 6.05, "患者単位5-fold CV：全患者が1回ずつvalidationになる",
              ha="center", fontsize=15, fontweight="bold", color="#1F3B63")
    colors = ["#D9EAF7", "#E2F0D9", "#FCE4D6", "#E4DFEC", "#FFF2CC"]
    for row in range(5):
        y = 4.95 - row * 0.9
        axis.text(0.25, y + 0.3, f"学習{row + 1}", fontsize=10, va="center")
        for col in range(5):
            x = 1.25 + col * 1.75
            is_validation = row == col
            box = FancyBboxPatch(
                (x, y), 1.5, 0.6,
                boxstyle="round,pad=0.02",
                edgecolor="#C00000" if is_validation else "#548B54",
                facecolor="#F4CCCC" if is_validation else colors[col],
                linewidth=2 if is_validation else 1,
            )
            axis.add_patch(box)
            axis.text(x + 0.75, y + 0.3,
                      f"Fold {col + 1}\n{'検証' if is_validation else '学習'}",
                      ha="center", va="center", fontsize=8.5,
                      fontweight="bold" if is_validation else "normal")
    axis.text(5.5, 0.25,
              "5つの検証予測を結合すると、1,154人全員のOOF（out-of-fold）予測になる",
              ha="center", fontsize=10.5, color="#555555")
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)


def create_model_figure(path: Path) -> None:
    plt.rcParams["font.family"] = FONT_JP
    fig, axes = plt.subplots(1, 2, figsize=(11, 4.4))
    for axis in axes:
        axis.set_xlim(0, 10)
        axis.set_ylim(0, 7)
        axis.axis("off")
    axes[0].set_title("CNN（ResNet18）のイメージ", fontsize=14, fontweight="bold", color="#C55A11")
    axes[1].set_title("ViT（DeiT-small）のイメージ", fontsize=14, fontweight="bold", color="#2E74B5")
    for row in range(4):
        for col in range(5):
            axes[0].add_patch(plt.Rectangle((1.2 + col * 1.3, 1.2 + row * 1.1), 1.0, 0.8,
                                            facecolor="#FCE4D6", edgecolor="#C55A11"))
    axes[0].add_patch(plt.Rectangle((2.5, 2.3), 2.3, 1.9, fill=False,
                                    edgecolor="#C00000", linewidth=3))
    axes[0].text(5, 0.45, "近くの模様を積み重ねて特徴を捉える\n『拡大鏡を少しずつ動かす』",
                 ha="center", fontsize=10)

    centers = []
    for row in range(4):
        for col in range(5):
            x, y = 1.2 + col * 1.3, 1.2 + row * 1.1
            axes[1].add_patch(plt.Rectangle((x, y), 1.0, 0.8,
                                            facecolor="#D9EAF7", edgecolor="#2E74B5"))
            centers.append((x + 0.5, y + 0.4))
    selected = [centers[1], centers[8], centers[12], centers[18]]
    for i, start in enumerate(selected):
        for end in selected[i + 1:]:
            axes[1].plot([start[0], end[0]], [start[1], end[1]], color="#8064A2", alpha=0.55)
    axes[1].text(5, 0.45, "画像を小領域に分け、離れた場所の関係も見る\n『地図上の地点を結び付ける』",
                 ha="center", fontsize=10)
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)


def add_title(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(36)
    paragraph.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run("MRI×AI研究\n初学者ガイドブック")
    apply_font(run, size=23, bold=True, color=NAVY)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("大脳白質病変grade分類とCNN・ViT比較を理解するために")
    apply_font(run, size=14, bold=True, color=BLUE)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(18)
    paragraph.paragraph_format.space_after = Pt(18)
    run = paragraph.add_run("医学・機械学習・統計の用語を、例えと図から段階的に学ぶ")
    apply_font(run, size=11, color=TEXT)

    for text in ("対象：AIやMRI研究を初めて読む学生", "研究記録の理解を助ける副読本"):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        apply_font(run, size=10.5, color=MUTED)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(24)
    run = paragraph.add_run("作成日：2026年7月17日")
    apply_font(run, size=10.5, color=MUTED)

    add_note(
        doc,
        "このガイドの役割",
        "本書は専門的な研究記録を置き換えるものではない。『なぜその実験が必要だったのか』『数字から何を言えて、何を言えないのか』を"
        "初学者が理解するための案内書である。最終的な実験条件・数値・限界は、研究全経過記録と照合する。",
        fill="E2F0D9",
    )
    doc.add_page_break()


def add_contents(doc: Document) -> None:
    add_heading(doc, "この本の使い方")
    add_para(doc, "最初に第1章と第2章を読めば研究の全体像が分かる。第3～10章は用語と方法、第11～13章は結果の読み方、"
                  "第14章以降は誤解しやすい点、今後の研究、用語集である。分からない単語は巻末の用語集へ戻ってよい。")
    add_bullets(doc, [
        ("短時間で把握したい：", "第1章、第11章、第13章を読む。"),
        ("方法を理解したい：", "第4～10章を順番に読む。"),
        ("論文・発表を準備したい：", "第11～16章を読み、研究全経過記録の数値表と照合する。"),
    ])
    add_heading(doc, "目次", 2)
    sections = [
        "1. 3分で分かる今回の研究", "2. まず押さえる5つのポイント", "3. MRIと白質病変の基礎",
        "4. 画像分類AIは何をしているか", "5. 最大の落とし穴：患者リーケージ", "6. なぜ3クラスにしたのか",
        "7. CNNとViTの違い", "8. 公平な比較条件", "9. 全axialと患者top-5",
        "10. 5-fold CV・OOF・bootstrap", "11. 評価指標の読み方", "12. 最終結果を一緒に読む",
        "13. 今回分かったこと・分からなかったこと", "14. よくある誤解", "15. 次に行う研究",
        "16. 研究記録と成果物の読み方", "17. よくある質問", "付録A. 用語集", "付録B. 1ページ要約",
    ]
    add_bullets(doc, sections)
    doc.add_page_break()


def build_document(asset_dir: Path) -> Document:
    pipeline_figure = asset_dir / "pipeline.png"
    split_figure = asset_dir / "split.png"
    cv_figure = asset_dir / "cv.png"
    model_figure = asset_dir / "models.png"
    create_pipeline_figure(pipeline_figure)
    create_split_figure(split_figure)
    create_cv_figure(cv_figure)
    create_model_figure(model_figure)

    doc = Document()
    setup_document(doc)
    add_title(doc)
    add_contents(doc)

    add_heading(doc, "1. 3分で分かる今回の研究")
    add_para(
        doc,
        "この研究の目的は、頭部MRI画像から大脳白質病変の程度を患者単位で自動分類し、"
        "代表的なCNNとVision Transformer（ViT）のどちらが有効かを公平に調べることである。"
        "対象は1,154人で、FLAIR（本研究ではFLと表記）とT1の全axial画像を使用した。",
    )
    add_figure(doc, pipeline_figure, "図1　今回の主解析を一枚で表した流れ")
    add_table(doc, ["問い", "今回の答え"], [
        ["AIは何を分類する？", "患者をgrade0 / grade1 / grade2+の3クラスへ分類"],
        ["画像は何を使う？", "FL＋T1の全axialスライス"],
        ["比較するモデルは？", "CNNのResNet18とViTのDeiT-small"],
        ["評価方法は？", "患者単位5-fold CV、各患者1回のOOF予測"],
        ["結論は？", "分類精度は同程度。ViTはmacro-AUCと確率品質で小さく優位"],
        ["重症検出はViTが上？", "grade2+の明確な優位は確認されなかった"],
    ])
    add_note(
        doc,
        "一言で表すと",
        "『ViTが何でも圧勝した』研究ではない。患者リーケージを除き、公平な5-fold評価を行った結果、"
        "CNNとViTの正解率はほぼ同じだったが、ViTは患者の順位付けと確率予測で小さい優位を示した研究である。",
        fill="FFF2CC",
    )

    add_heading(doc, "2. まず押さえる5つのポイント")
    add_bullets(doc, [
        ("① 評価単位は患者：", "1人に多数の画像があっても、独立した症例は1人である。"),
        ("② 同じ患者をtrainとvalidationへ混ぜない：", "混ぜると、未見患者への性能を測れない。"),
        ("③ 5クラスを3クラスへ統合：", "grade4が8人しかおらず、安定した個別分類が困難だった。"),
        ("④ 全axialを入力しtop-5で患者集約：", "手動の範囲選択を避け、モデルが有用と判断した5枚を平均した。"),
        ("⑤ 大差より小さく堅牢な差を重視：", "単一splitの大差ではなく、1,154人OOFで再現したAUC差を結論にした。"),
    ])
    add_note(
        doc,
        "研究で最も大切だったこと",
        "モデルを複雑にすることより、患者リーケージを見つけ、評価設計を修正したことの方が重要である。"
        "正しい評価がなければ、高いAccuracyにも意味がない。",
        fill="E2F0D9",
    )

    doc.add_page_break()
    add_heading(doc, "3. MRIと白質病変の基礎")
    add_heading(doc, "3.1 MRIとは", 2)
    add_para(
        doc,
        "MRIは磁場と電波を利用して体内を画像化する。CTとは原理が異なり、撮像条件を変えることで同じ脳を異なる見え方で観察できる。"
        "この『異なる見え方』をモダリティまたは撮像系列と呼ぶ。",
    )
    add_table(doc, ["撮像系列", "初学者向けの役割", "今回の扱い"], [
        ["FLAIR（FL）", "脳脊髄液を抑え、白質病変が明るく見えやすい", "主要な病変情報"],
        ["T1", "脳の解剖学的な形や組織境界を把握しやすい", "形態情報を補助"],
        ["T2", "水分の多い部分が明るく見えやすい", "今回の主解析には含めない"],
    ])
    add_heading(doc, "3.2 axialスライスとは", 2)
    add_para(
        doc,
        "axialは頭を上下方向に薄く切ったように見る断面である。1人の患者から複数の高さの画像が得られる。"
        "たとえると、1冊の本を1ページずつ撮影したものに近い。ページ数が多くても、本そのものは1冊である。",
    )
    add_heading(doc, "3.3 大脳白質病変とgrade", 2)
    add_para(
        doc,
        "大脳白質病変は脳の白質に見られる変化で、FLAIR画像では高信号として見えることが多い。"
        "本データには患者単位のwm grade 0～4が付与されている。gradeが大きいほど病変程度が高い区分として扱った。"
        "AIが直接医学的診断を確定するのではなく、与えられた教師gradeを予測する。",
    )
    add_note(
        doc,
        "重要",
        "ラベルは画像1枚ごとではなく患者全体に付いている。同じ患者の病変が写りにくいスライスにも同じgradeが付くため、"
        "スライス単位学習にはラベルノイズに似た難しさが生じる。",
    )

    add_heading(doc, "4. 画像分類AIは何をしているか")
    add_para(
        doc,
        "学習では、MRI画像と正解gradeの組を繰り返しモデルへ見せる。モデルは各クラスらしさを確率として出力し、"
        "正解に近づくよう内部パラメータを更新する。評価では学習に使っていない患者へ適用する。",
    )
    add_table(doc, ["用語", "役割", "学校の試験にたとえると"], [
        ["train", "モデルが重みを学習するデータ", "練習問題"],
        ["validation", "学習途中のモデル選択や性能確認", "模擬試験"],
        ["test / external validation", "設定固定後の最終評価", "本試験"],
        ["checkpoint", "あるepoch時点のモデル保存", "途中時点の答案・実力記録"],
        ["inference", "学習済みモデルで予測すること", "本番で問題を解く"],
    ])
    add_note(
        doc,
        "過学習とは",
        "練習問題を暗記してtrain性能だけが高くなり、新しい問題では正解できない状態である。"
        "今回の初期ViTでもtrain Accuracyが約0.99まで上昇した一方、validation lossが悪化した。",
        fill="FCE4D6",
    )

    doc.add_page_break()
    add_heading(doc, "5. 最大の落とし穴：患者リーケージ")
    add_para(
        doc,
        "初期研究では画像をランダムにtrainと評価側へ分けていた。しかし同じ患者の別スライスが両方へ入ると、"
        "モデルは病変の一般則だけでなく、その患者固有の形や撮像特徴を覚えられる。これが患者リーケージである。",
    )
    add_figure(doc, split_figure, "図2　画像単位splitと患者単位splitの違い")
    add_table(doc, ["分割", "何が混ざるか", "測っている性能"], [
        ["画像単位split", "同一患者の別画像がtrainと評価側へ入る", "見たことがある患者を含む性能"],
        ["患者単位split", "患者の全画像を片側だけへ置く", "未見患者への汎化性能"],
    ])
    add_para(
        doc,
        "旧分割ではvalidation患者1,148人中1,146人がtrainにも存在し、ほぼ100%の患者リーケージだった。"
        "そこで得た約0.92のAccuracyは、現在の主結果と比較できない参考値である。",
    )
    add_note(
        doc,
        "覚え方",
        "『同じ人の写真を練習問題と試験問題の両方へ入れない』。医用画像では、画像数ではなく患者IDを基準に分ける。",
        fill="FFF2CC",
    )

    add_heading(doc, "6. なぜ3クラスにしたのか")
    add_table(doc, ["元grade", "患者数", "3クラスでの扱い"], [
        ["grade0", 557, "grade0"],
        ["grade1", 444, "grade1"],
        ["grade2", 106, "grade2+"],
        ["grade3", 39, "grade2+"],
        ["grade4", 8, "grade2+"],
    ])
    add_para(
        doc,
        "grade4は全体で8人しかいない。同一患者に多数の画像があっても独立患者数は増えないため、grade4だけを安定して学習・評価することは難しい。"
        "そこでgrade2・3・4をgrade2+へ統合し、grade0 / grade1 / grade2+の3クラスとした。",
    )
    add_note(
        doc,
        "統合はごまかしか",
        "目的と症例数に基づいて事前に定義し、元grade別の結果も併記すれば合理的である。"
        "ただし3クラスモデルはgrade2・3・4を互いに区別できず、『grade4を分類した』とは言えない。",
    )
    add_heading(doc, "6.1 クラス不均衡", 2)
    add_para(
        doc,
        "クラス不均衡とは、クラスごとの患者数が偏ることである。少数クラスへ大きな重みを付ければ必ず改善するわけではない。"
        "今回、samplerと強いclass weightを重ねると少数クラスを過剰予測し、穏やかなclass-balanced loss単独も改善しなかった。"
        "最も効果が大きかったのは、重み調整より3クラス化だった。",
    )

    doc.add_page_break()
    add_heading(doc, "7. CNNとViTの違い")
    add_figure(doc, model_figure, "図3　CNNとViTの考え方を単純化したイメージ")
    add_table(doc, ["観点", "CNN（ResNet18）", "ViT（DeiT-small）"], [
        ["基本動作", "近傍の模様を畳み込みで積み重ねる", "画像patch間の関係をattentionで扱う"],
        ["得意と期待される点", "局所模様、効率、少量データでの安定性", "離れた領域の関係、全体構造"],
        ["今回のモデル", "ResNet18", "deit_small_patch16_224"],
        ["注意点", "条件が違えば比較は不公平", "大規模モデルは少数患者で過学習しやすい"],
    ])
    add_heading(doc, "7.1 旧ViTが不利だった理由", 2)
    add_bullets(doc, [
        "80px画像を16px patchに分け、5×5=25トークンしかなかった。",
        "vit_baseは約86M parametersで、1,154患者規模に対して大きかった。",
        "CNN再現用のRMSprop・augmentationなしをそのまま使用した。",
    ])
    add_para(
        doc,
        "そこで224px、約21.7M parametersのDeiT-small、AdamW、augmentation、weight decay、dropout、drop path、"
        "label smoothing、cosine schedulerへ変更した。これはViTを優遇するためではなく、ViTが通常必要とする条件で評価するためである。",
    )

    add_heading(doc, "8. 公平な比較条件")
    add_para(doc, "モデル以外の条件が違えば、性能差がCNN/ViTの構造によるものか判断できない。そこで以下を固定した。")
    add_table(doc, ["条件", "固定内容"], [
        ["患者・画像", "同一患者split、FL＋T1、全axial"],
        ["入力", "224×224px"],
        ["学習", "30epoch、batch size 16"],
        ["optimizer", "AdamW、learning rate 3e-5、weight decay 0.1"],
        ["その他", "cosine scheduler、label smoothing 0.1、同一augmentation"],
        ["checkpoint", "validation lossが最小のモデル"],
        ["患者集約", "信頼度上位5枚の確率平均"],
    ])
    add_note(
        doc,
        "公平＝すべて完全に同じ、ではない",
        "ViT固有のdropout/drop pathは構造に合わせた正則化で、CNNには同じ形では存在しない。"
        "公平性とは、目的と評価を揃え、各モデルに妥当な学習条件を与え、その違いを明記することである。",
    )

    doc.add_page_break()
    add_heading(doc, "9. 全axialと患者top-5")
    add_heading(doc, "9.1 なぜ全axialか", 2)
    add_para(
        doc,
        "axial 9～15へ手動制限すると重症Recallが改善した参考実験がある。しかし装置・撮像範囲・頭部位置が変わると、"
        "同じ番号が同じ解剖位置とは限らない。境界選択にも研究者の恣意性が入る。そこで主解析は全axialに固定した。",
    )
    add_heading(doc, "9.2 top-5とは", 2)
    add_para(
        doc,
        "モデルは患者の全スライスへ3クラス確率を出す。各スライスの最大確率を『信頼度』とし、上位5枚の確率を平均して患者予測を作る。"
        "病変が写らない画像を含む全平均より重要信号を残し、1枚だけを選ぶmax confidenceより偶然の極端値へ依存しにくい。",
    )
    add_table(doc, ["集約方法", "たとえ", "長所・短所"], [
        ["mean", "本の全ページを同じ重みで読む", "安定だが重要ページが薄まる"],
        ["max confidence", "最も強い1ページだけ読む", "病変を拾えるが偶然の誤信頼に弱い"],
        ["top-5", "重要そうな5ページを読み合わせる", "信号と安定性の折衷"],
    ])
    add_note(
        doc,
        "top-5はViT内部attentionではない",
        "各スライスのsoftmax信頼度による評価時の後処理であり、スライス間関係を学習した患者Transformerではない。",
        fill="FFF2CC",
    )
    add_heading(doc, "9.3 モデルが選んだ画像", 2)
    add_table(doc, ["全1,154患者", "CNN", "ViT"], [
        ["top-5中のFL割合", "73.1%", "74.1%"],
        ["axial 9～15の割合", "41.7%", "54.5%"],
        ["axial中央値", "13", "13"],
    ])
    add_para(
        doc,
        "全axialを入力しても、両モデルはFLを多く選び、ViTは中央付近をより多く選んだ。"
        "病変情報との整合性を示す探索的所見だが、医学的な因果説明ではない。",
    )

    add_heading(doc, "10. 5-fold CV・OOF・bootstrap")
    add_heading(doc, "10.1 5-fold cross-validation", 2)
    add_figure(doc, cv_figure, "図4　患者単位5-fold cross-validation")
    add_para(
        doc,
        "1,154患者をおおむね5等分し、4部分で学習、残り1部分で評価する操作を5回行う。"
        "患者はvalidationに1回だけ入り、trainとvalidationの患者重複は各foldで0である。",
    )
    add_heading(doc, "10.2 OOF予測", 2)
    add_para(
        doc,
        "OOFはout-of-foldの略である。各患者について、その患者を学習に使っていないfoldモデルの予測だけを集める。"
        "5つのvalidation予測を結合すると、1,154人全員の公平な予測一覧になる。",
    )
    add_heading(doc, "10.3 paired bootstrap", 2)
    add_para(
        doc,
        "同じ患者にCNNとViTを適用しているため、比較は対応あり（paired）である。患者を復元抽出して10,000個の仮想標本を作り、"
        "毎回ViT－CNN差を計算した。差の中央95%範囲を95%信頼区間とした。クラス比が崩れないよう各クラス内で抽出した。",
    )
    add_note(
        doc,
        "95%信頼区間の読み方",
        "モデル差の区間が0をまたぐ場合、今回のデータだけでは『どちらが上か』を確定的に言いにくい。"
        "0をまたがない場合でも、差の大きさ・臨床的意味・外部再現性を別に考える必要がある。",
    )

    doc.add_page_break()
    add_heading(doc, "11. 評価指標の読み方")
    add_table(doc, ["指標", "何を見るか", "今回の注意"], [
        ["Accuracy", "全患者のうち正解した割合", "多数クラスの影響を受ける"],
        ["Recall", "そのクラス患者をどれだけ拾えたか", "grade2+を見逃さないか確認"],
        ["macro-F1", "クラスごとのF1を同じ重みで平均", "少数クラスも同等に扱う"],
        ["ROC-AUC", "閾値を動かした順位付け性能", "高くても固定閾値のAccuracy差とは限らない"],
        ["PR-AUC/AP", "precisionとrecallの両立", "少数のgrade2+で特に重要"],
        ["NLL", "正解クラスへ置いた確率の質", "自信を持った誤りを強く罰する"],
        ["Brier score", "予測確率と正解の二乗誤差", "低いほど確率全体が良い"],
        ["ECE", "自信度と実際の正解率のずれ", "低いほど較正が良い"],
    ], font_size=8.0)
    add_heading(doc, "11.1 AUCを直感的に理解する", 2)
    add_para(
        doc,
        "たとえばgrade2+患者と非grade2+患者を1人ずつ無作為に選んだとき、モデルがgrade2+患者へより高いgrade2+確率を付ける能力がAUCに対応する。"
        "AUC=0.5は偶然程度、1.0は完全な順位付けである。AUCは『どの閾値でクラスを決めたか』とは別の性質を見る。",
    )
    add_heading(doc, "11.2 calibrationを直感的に理解する", 2)
    add_para(
        doc,
        "モデルが『80%の自信』とした患者群の約80%が実際に正解なら、確率はよく較正されている。"
        "Accuracyが同じでも、確率が信頼できるモデルはリスク順位付けや二次判定へ使いやすい。",
    )
    add_heading(doc, "11.3 McNemar検定", 2)
    add_para(
        doc,
        "Accuracyの比較では、ViTだけ正解した患者数とCNNだけ正解した患者数を比較する。"
        "今回のOOFでは92人対84人で、exact McNemar p=0.5979だった。したがってAccuracy優位は示されなかった。",
    )

    add_heading(doc, "12. 最終結果を一緒に読む")
    add_heading(doc, "12.1 全1,154患者の主結果", 2)
    add_table(doc, ["指標", "CNN", "ViT", "ViT－CNN 95% CI", "読み方"], [
        ["Accuracy", "0.6282", "0.6352", "－0.0156～0.0295", "差は確定しない"],
        ["macro-F1", "0.6001", "0.6134", "－0.0135～0.0399", "差は確定しない"],
        ["macro-AUC", "0.7891", "0.8100", "0.0062～0.0357", "ViTが小さく優位"],
        ["balanced accuracy", "0.5995", "0.6132", "－0.0112～0.0386", "差は確定しない"],
        ["grade2+ Recall", "0.5882", "0.6144", "－0.0261～0.0784", "差は確定しない"],
    ], font_size=7.8)
    add_note(
        doc,
        "ここが研究の中心",
        "ViTはAccuracyで約0.7ポイント、macro-AUCで約2.1ポイント高かった。"
        "Accuracy差の信頼区間は0をまたぐが、macro-AUC差は0をまたがない。したがって『正解率で圧勝』ではなく、"
        "『順位付け性能に小さい優位』と表現する。",
        fill="E2F0D9",
    )
    add_figure(doc, DETAIL_FIGURES / "class_roc_curves.png", "図5　クラス別ROC曲線：曲線が左上に近いほど良い")

    add_heading(doc, "12.2 どのクラスで差が出たか", 2)
    add_table(doc, ["クラス", "CNN ROC-AUC", "ViT ROC-AUC", "差の95% CI"], [
        ["grade0", "0.7845", "0.8122", "0.0119～0.0437"],
        ["grade1", "0.6963", "0.7144", "－0.0037～0.0401"],
        ["grade2+", "0.8866", "0.9035", "－0.0048～0.0400"],
    ])
    add_para(
        doc,
        "すべてViTの点推定が高いが、差が確定的なのはgrade0だけだった。grade2+のPR-AUCはCNN 0.7097、ViT 0.7104でほぼ同じだった。"
        "よって『ViTは重症病変の検出に優れる』とは結論できない。",
    )
    add_heading(doc, "12.3 確率予測の品質", 2)
    add_table(doc, ["低いほど良い指標", "CNN", "ViT", "ViT－CNN 95% CI"], [
        ["NLL", "0.7894", "0.7532", "－0.0590～－0.0140"],
        ["Multiclass Brier", "0.4864", "0.4627", "－0.0398～－0.0081"],
        ["Top-label ECE", "0.0719", "0.0579", "－0.0345～0.0096"],
    ])
    add_para(doc, "ViTはNLLとBrier scoreで良好だった。AUCだけでなく、正解クラスへ割り当てる確率の質にも小さい優位がある。ECE差は確定しなかった。")
    add_figure(doc, DETAIL_FIGURES / "top_label_reliability.png", "図6　自信度と実際の正解率：点が対角線に近いほど較正が良い", width=5.6)

    add_heading(doc, "12.4 どの患者で間違えたか", 2)
    add_table(doc, ["結果", "患者数", "意味"], [
        ["両モデル正解", 641, "どちらでも分類できた"],
        ["ViTのみ正解", 92, "ViTが補えた"],
        ["CNNのみ正解", 84, "CNNが補えた"],
        ["両モデル不正解", 337, "データ・ラベル・モデル共通の難しさ"],
    ])
    add_para(
        doc,
        "grade1患者444人のうち233人を両モデルが誤分類した。grade1はgrade0との境界が曖昧で、今回の主要な改善対象である。"
        "一方だけが正解した患者は176人いるためensembleの可能性はあるが、実際に学習・評価しなければ性能向上とは言えない。",
    )
    add_figure(doc, DETAIL_FIGURES / "confusion_matrices.png", "図7　混同行列：縦が正解、横がモデル予測")
    add_figure(doc, DETAIL_FIGURES / "error_overlap.png", "図8　同じ患者に対するCNNとViTの正誤関係", width=5.5)

    doc.add_page_break()
    add_heading(doc, "13. 今回分かったこと・分からなかったこと")
    add_table(doc, ["今回言えること", "今回だけでは言えないこと"], [
        ["患者リーケージを除いた1,154人OOF性能", "別施設でも同じ性能が出ること"],
        ["CNNとViTのAccuracy/F1は同程度", "ViTが分類精度でCNNを圧倒すること"],
        ["ViTのmacro-AUCに小さい優位", "すべてのクラスでViTが確実に上であること"],
        ["ViTのNLL・Brierが良好", "確率を臨床判断へそのまま使用できること"],
        ["grade0のAUC/AP差が明確", "grade2+検出でViTが優れること"],
        ["grade4は両モデル8/8", "grade4一般化性能が100%であること"],
    ], font_size=8.0)
    add_note(
        doc,
        "最終的な研究表現",
        "全axial・患者単位5-fold CVにおいて、DeiT-smallはResNet18と同程度の分類性能を示し、"
        "macro-AUCと確率予測品質では小さいが再現性のある優位性を示した。"
        "ただしgrade2+の検出優位性は確認されなかった。",
        fill="E2F0D9",
    )

    add_heading(doc, "14. よくある誤解")
    add_table(doc, ["誤解", "正しい理解"], [
        ["画像が53,194枚あるので標本数も53,194", "独立性を考えると主な標本数は患者1,154人"],
        ["Accuracyが高いモデルが常に良い", "不均衡ではF1、Recall、AUC、較正も必要"],
        ["AUCが高いのでAccuracyも必ず高い", "AUCは順位、Accuracyは固定されたクラス決定"],
        ["信頼区間が0をまたぐ＝完全に同じ", "差を確定する証拠が不足しているという意味"],
        ["p値が大きい＝差が存在しない証明", "今回の標本では差を検出できなかっただけ"],
        ["attention/top-5画像が医学的原因", "モデルが重視した候補であり、専門医確認が必要"],
        ["grade4 8/8＝grade4分類100%", "3クラスでgrade2+として検出しただけで患者も8人"],
        ["単一splitでViTが上なので結論", "5-foldで差が縮小したため主解析はOOF結果"],
    ], font_size=8.0)

    add_heading(doc, "15. 次に行う研究")
    add_heading(doc, "15.1 最優先：外部検証", 2)
    add_para(
        doc,
        "別施設または別撮像期間の患者へ、現在の全axial・top-5・モデル・前処理を変更せず適用する。"
        "施設、装置、画質が変わってもViTのAUC・確率品質が維持されるかを確認する。",
    )
    add_heading(doc, "15.2 外部データがない場合", 2)
    add_bullets(doc, [
        "異なる患者分割seedを2つ追加し、合計3反復のrepeated 5-fold CVを行う。",
        "学習条件を変更せず、macro-AUC差の方向と大きさを確認する。",
        "grade1の両モデル誤分類例を、予測結果を隠して専門医が再確認する。",
    ])
    add_heading(doc, "15.3 将来のモデル", 2)
    add_para(
        doc,
        "手動でaxial範囲を決めず、全スライスのembeddingを患者bagとして入力するattention/MILモデルが候補である。"
        "CNNとViTへ同じ患者aggregatorを適用し、nested CVまたは新しい外部データで公平に評価する。",
    )
    add_note(
        doc,
        "今は変更しないこと",
        "現在のOOF結果を見ながらtop-k、閾値、augmentationを調整すると、1,154人が再び開発データとなる。"
        "主解析条件は固定し、新しい工夫はinner validationまたは別データで選ぶ。",
        fill="FCE4D6",
    )

    doc.add_page_break()
    add_heading(doc, "16. 研究記録と成果物の読み方")
    add_para(
        doc,
        "本ガイドで考え方を理解した後、研究全経過記録で実験順序、設定、否定結果、最終数値を確認する。"
        "主解析は『全axial・患者単位5-fold CV』と『OOF予測の詳細解析』の節である。",
    )
    add_table(doc, ["成果物", "役割"], [
        ["MRI_ViT研究_全経過記録_20260717.docx", "研究の時系列、全実験、結果、限界を記録"],
        ["MRI_ViT研究_初学者ガイドブック_20260717.docx", "用語・方法・結果の読み方を初学者向けに説明"],
        ["oof_summary", "1,154患者の主指標とpaired bootstrap"],
        ["oof_detailed_analysis", "クラス別ROC/PR、較正、エラー、元grade、図表"],
        ["generated_cv5_all_axial_3class/PROTOCOL.md", "5-fold前に固定した条件"],
    ])
    add_heading(doc, "16.1 結果を確認する順番", 2)
    add_bullets(doc, [
        "まず患者数・クラス数・split方法を確認する。",
        "次にAccuracyだけでなくmacro-F1、AUC、各クラスRecallを見る。",
        "ViT－CNN差の95%信頼区間が0をまたぐか確認する。",
        "混同行列で、どのクラス間の誤りが多いかを見る。",
        "最後にselection bias、症例数、外部検証の有無を確認する。",
    ])

    add_heading(doc, "17. よくある質問")
    faqs = [
        ("Q1. なぜ2D画像を使うのですか？", "既存データとモデル資源を活用しやすく、各スライスの特徴をImageNet事前学習モデルで扱えるためである。ただし患者内の3D関係を直接学習していない限界がある。"),
        ("Q2. なぜFLだけにしないのですか？", "FLが多く選択されたが、T1が補助情報を持つ可能性がある。FLのみとの公平なablationは今後の検討対象である。"),
        ("Q3. top-5の5は最適ですか？", "単一validationでtop-3やtop-5を比較した後に固定した候補であり、普遍的な最適値ではない。現在のOOFで再選択してはいけない。"),
        ("Q4. ViTの優位は臨床的に大きいですか？", "macro-AUC差は約0.021で小さい。統計的に0をまたがなかったことと、臨床的に十分大きいことは別である。"),
        ("Q5. なぜgrade1が難しいのですか？", "病変なしと軽度の境界で画像所見が連続的であり、患者内スライス差やラベル判定の揺らぎが影響する可能性がある。専門医監査が必要である。"),
        ("Q6. 研究は完成ですか？", "内部5-fold主解析は完成した。一般化を示すには外部施設検証、または少なくともrepeated CVが必要である。"),
    ]
    for question, answer in faqs:
        add_note(doc, question, answer, fill="F3F6FA")

    doc.add_page_break()
    add_heading(doc, "付録A. 用語集")
    add_table(doc, ["用語", "やさしい説明"], [
        ["axial", "頭部を上下方向に輪切りにした断面"],
        ["FLAIR / FL", "白質病変が比較的見やすいMRI撮像系列"],
        ["T1", "脳の解剖学的構造を見やすいMRI撮像系列"],
        ["class", "AIが分類する区分。今回はgrade0 / grade1 / grade2+"],
        ["label", "教師として与える正解"],
        ["patient split", "同じ患者の全画像を同じ側へまとめる分割"],
        ["patient leakage", "同じ患者が学習側と評価側へ混在すること"],
        ["epoch", "全trainデータを一通り学習する単位"],
        ["augmentation", "回転や明るさ変更で学習画像へ変化を加えること"],
        ["overfitting", "学習データを覚えすぎ、新しいデータへ汎化しないこと"],
        ["CNN", "近傍の画像特徴を段階的に捉えるニューラルネットワーク"],
        ["ViT", "画像patchをtokenとしてattentionで関係付けるモデル"],
        ["softmax probability", "各クラスらしさを合計1の確率へ変換した値"],
        ["pooling / aggregation", "複数スライス予測を患者予測へまとめる処理"],
        ["top-5", "信頼度上位5枚のクラス確率を平均する集約"],
        ["cross-validation", "学習・評価部分を入れ替えて複数回評価する方法"],
        ["OOF", "その患者を学習に使っていないfoldモデルの予測"],
        ["bootstrap", "標本を復元抽出して統計的不確実性を推定する方法"],
        ["95% CI", "推定値の不確実性を表す95%信頼区間"],
        ["macro-F1", "各クラスを同じ重みで扱ったF1平均"],
        ["ROC-AUC", "閾値に依存しない順位付け性能"],
        ["PR-AUC / AP", "少数陽性クラスでprecisionとrecallを評価する指標"],
        ["calibration", "予測自信度と実際の正解率の一致度"],
        ["NLL", "正解へ低い確率を付けた場合、特に自信ある誤りを罰する指標"],
        ["Brier score", "予測確率と正解の二乗誤差"],
        ["ECE", "自信度と正解率のずれをbinごとにまとめた値"],
        ["MIL", "複数画像を患者bagとして患者ラベルで学習する方法"],
    ], font_size=7.8)

    add_heading(doc, "付録B. 1ページ要約")
    add_table(doc, ["項目", "固定内容・結果"], [
        ["研究目的", "MRIから患者wm gradeを分類し、CNNとViTを公平比較"],
        ["対象", "1,154患者、53,194画像、FL＋T1、全axial"],
        ["クラス", "grade0=557、grade1=444、grade2+=153"],
        ["モデル", "ResNet18 vs DeiT-small 224"],
        ["評価", "患者単位5-fold OOF、best-loss、top-5"],
        ["CNN", "Acc 0.6282 / F1 0.6001 / AUC 0.7891"],
        ["ViT", "Acc 0.6352 / F1 0.6134 / AUC 0.8100"],
        ["確定的な全体差", "macro-AUC +0.0209［95% CI 0.0062～0.0357］"],
        ["AUC差の主な由来", "grade0 ROC-AUCとPR-AUC"],
        ["grade2+", "ROC/PR差は確定せず、ViT優位は未証明"],
        ["確率品質", "ViTのNLL・Brierが良好"],
        ["主要課題", "grade1：444人中233人を両モデルが誤分類"],
        ["結論", "分類性能は同程度。ViTはAUC・確率品質で小さく優位"],
        ["次の検証", "外部施設検証。なければrepeated patient-level CV"],
    ], font_size=8.0)
    add_note(
        doc,
        "最後に",
        "この研究を理解する鍵は、モデル名よりも『患者単位で正しく分けたか』『どの単位で評価したか』『差の不確実性を示したか』である。"
        "高い数字だけを選ばず、再現しなかった結果や限界も記録することが科学的な研究につながる。",
        fill="E2F0D9",
    )
    return doc


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="mri_vit_guide_") as temp_dir:
        document = build_document(Path(temp_dir))
        document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
