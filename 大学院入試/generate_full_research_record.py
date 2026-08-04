# -*- coding: utf-8 -*-
"""MRI/ViT研究プロジェクトの全経過記録をWord形式で生成する。"""

from __future__ import annotations

import csv
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
MRI_ROOT = ROOT / "mri-vit-classification"
OUTPUT_PATH = Path(__file__).resolve().parent / "MRI_ViT研究_全経過記録_20260717.docx"

FONT_JP = "Yu Gothic"
COLOR_NAVY = RGBColor(0x1F, 0x3B, 0x63)
COLOR_BLUE = RGBColor(0x2E, 0x74, 0xB5)
COLOR_TEXT = RGBColor(0x22, 0x22, 0x22)
COLOR_MUTED = RGBColor(0x66, 0x66, 0x66)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def apply_font(run, size: float | None = None, bold: bool | None = None,
               color: RGBColor | None = None, italic: bool | None = None) -> None:
    run.font.name = FONT_JP
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("eastAsia", "ascii", "hAnsi"):
        rfonts.set(qn(f"w:{attr}"), FONT_JP)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if italic is not None:
        run.font.italic = italic


def set_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT_JP
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = COLOR_TEXT
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_JP)

    heading_settings = {
        "Title": (22, COLOR_NAVY),
        "Heading 1": (15, COLOR_NAVY),
        "Heading 2": (12.5, COLOR_BLUE),
        "Heading 3": (11, COLOR_BLUE),
    }
    for name, (size, color) in heading_settings.items():
        style = doc.styles[name]
        style.font.name = FONT_JP
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_JP)

    for name in ("List Bullet", "List Bullet 2", "List Number", "Caption"):
        style = doc.styles[name]
        style.font.name = FONT_JP
        style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_JP)


def configure_page(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Mm(18)
        section.bottom_margin = Mm(16)
        section.left_margin = Mm(18)
        section.right_margin = Mm(18)
        section.header_distance = Mm(8)
        section.footer_distance = Mm(8)

        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header.add_run("MRI画像による大脳白質病変グレード分類 ― 研究全経過記録")
        apply_font(run, size=8.5, color=COLOR_MUTED)

        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("- ")
        apply_font(run, size=9, color=COLOR_MUTED)
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.append(begin)
        run._r.append(instr)
        run._r.append(end)
        tail = footer.add_run(" -")
        apply_font(tail, size=9, color=COLOR_MUTED)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(38)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("MRI画像による大脳白質病変\nグレード分類")
    apply_font(run, size=22, bold=True, color=COLOR_NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run("研究プロジェクト 全経過記録")
    apply_font(run, size=17, bold=True, color=COLOR_BLUE)

    for text in (
        "CNNによる初期研究から、患者リーケージの発見、患者単位再評価、",
        "ViT最適化、患者集約、公平比較、患者単位5-fold CVまで",
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        apply_font(run, size=11, color=COLOR_TEXT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(26)
    run = p.add_run("作成者：竹村 典晃")
    apply_font(run, size=11, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("記録基準日：2026年7月17日")
    apply_font(run, size=10.5, color=COLOR_MUTED)

    add_note(
        doc,
        "本書の位置付け",
        "本書は、これまでに実施した実験、実装変更、失敗・否定結果、評価上の問題、"
        "得られた知見を一つにまとめた研究ログである。主解析は全1,154患者を各1回だけvalidationとする"
        "患者単位5-fold out-of-fold評価である。これ以前の289患者単一validation結果は、設定選択を含む探索的結果として区別する。",
        fill="EAF2F8",
    )
    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(level=level)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(10 if level == 1 else 7)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    apply_font(
        run,
        size=15 if level == 1 else (12.5 if level == 2 else 11),
        bold=True,
        color=COLOR_NAVY if level == 1 else COLOR_BLUE,
    )


def add_para(doc: Document, text: str, bold: bool = False,
             color: RGBColor = COLOR_TEXT, space_after: float = 5) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.18
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    apply_font(run, size=10.5, bold=bold, color=color)


def add_bullets(doc: Document, items: Iterable[str | tuple[str, str]], level: int = 1) -> None:
    style = "List Bullet" if level == 1 else "List Bullet 2"
    for item in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.12
        if isinstance(item, tuple):
            head, body = item
            r1 = p.add_run(head)
            apply_font(r1, size=10.5, bold=True)
            r2 = p.add_run(body)
            apply_font(r2, size=10.5)
        else:
            run = p.add_run(item)
            apply_font(run, size=10.5)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[object]],
              font_size: float = 8.7) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    repeat_header(table.rows[0])

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade_cell(cell, "2E74B5")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(header))
        apply_font(run, size=font_size, bold=True, color=COLOR_WHITE)

    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for col_idx, value in enumerate(row):
            cell = cells[col_idx]
            cell.text = ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_idx % 2 == 1:
                shade_cell(cell, "F3F6FA")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(value))
            apply_font(run, size=font_size)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)


def add_note(doc: Document, title: str, body: str, fill: str = "FFF2CC") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    apply_font(r, size=10.5, bold=True, color=COLOR_NAVY)
    p = cell.add_paragraph()
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(body)
    apply_font(r, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    apply_font(run, size=9, color=COLOR_MUTED, italic=True)


def add_figure(doc: Document, path: Path, caption: str, width: float = 6.5) -> None:
    if not path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def read_cnn_progress() -> dict[str, object] | None:
    path = MRI_ROOT / "outputs/repro_cnn_all_axial_patient_split_3class_resnet18_224_reg/logs/resnet18_epoch_log.csv"
    if not path.exists():
        return None
    with path.open(encoding="utf-8") as f:
        rows = list(csv.DictReader(f))
    if not rows:
        return None

    def best(key: str, mode: str = "max") -> tuple[int, float]:
        pairs = [(int(row["epoch"]), float(row[key])) for row in rows]
        return (max if mode == "max" else min)(pairs, key=lambda item: item[1])

    latest = rows[-1]
    return {
        "latest_epoch": int(latest["epoch"]),
        "latest_train_acc": float(latest["train_acc"]),
        "latest_val_acc": float(latest["val_acc"]),
        "latest_val_f1": float(latest["val_f1"]),
        "latest_val_auc": float(latest["val_roc_auc"]),
        "best_loss": best("val_loss", "min"),
        "best_acc": best("val_acc"),
        "best_f1": best("val_f1"),
        "best_auc": best("val_roc_auc"),
    }


def add_contents(doc: Document) -> None:
    add_heading(doc, "目次", 1)
    sections = [
        "1. 要約（現在地）", "2. 研究背景と目的", "3. データとラベル構造",
        "4. 初期CNN研究と高精度結果", "5. 患者リーケージの発見と再評価",
        "6. 患者単位split・5クラス実験", "7. 3クラス化",
        "8. クラス不均衡対策", "9. axial 9–15実験", "10. ViT向け条件への転換",
        "11. 正則化DeiT-small実験", "12. 患者集約", "13. top-k集約と選択画像分析",
        "14. 元グレード別サブグループ解析", "15. 公平比較CNNと単一split比較",
        "16. 全axial・患者単位5-fold CV", "17. OOF予測の詳細解析",
        "18. 実装した機能", "19. 得られた知見", "20. 失敗・否定結果",
        "21. 評価上の注意と限界", "22. 今後の計画", "23. 結論",
        "付録A. 実験結果一覧", "付録B. 主要ファイル", "付録C. 用語と数式",
    ]
    add_bullets(doc, sections)
    doc.add_page_break()


def build_document() -> Document:
    doc = Document()
    set_styles(doc)
    configure_page(doc)
    add_title(doc)
    add_contents(doc)

    add_heading(doc, "1. 要約（現在地）")
    add_para(
        doc,
        "本研究では、MRI画像から大脳白質病変（white matter lesion）のグレードを自動判別することを目的として、"
        "CNN（ResNet18）とVision Transformer（ViT/DeiT）を比較してきた。初期研究では約0.92の高精度を得たが、"
        "後に画像単位splitによる患者リーケージが判明した。その後、患者単位splitへ移行し、患者間汎化を厳密に評価した。",
    )
    add_bullets(doc, [
        ("最重要の修正：", "train/validationを患者単位で分離し、患者重複を0にした。"),
        ("ラベル構造：", "ラベルは患者単位であり、同一患者の約44～50枚のFL/T1スライスが同じgradeを持つ。"),
        ("クラス設計：", "極端な少数クラスを考慮し、grade0 / grade1 / grade2+（2・3・4統合）の3クラスへ変更した。"),
        ("主解析：", "全axial・患者単位5-fold CVにより、1,154患者全員のout-of-fold予測を取得した。"),
        ("分類性能：", "CNNとViTのAccuracy・macro-F1に明確な差はなく、ViTのmacro-AUCのみ小さいが再現性のある優位を示した。"),
        ("確率予測：", "ViTはNLLとmulticlass Brier scoreでも有意に良好だった。"),
    ])
    add_note(
        doc,
        "固定した主結論",
        "全axial・患者単位5-fold CVにおいて、DeiT-smallはResNet18と同程度の分類性能を示した。"
        "一方、macro-AUCはViTが0.0209高く、paired bootstrap 95%信頼区間は0.0062～0.0357で0をまたがなかった。"
        "この差は特にgrade0の識別で明確だったが、grade2+の検出優位性は確認されなかった。",
    )

    add_heading(doc, "2. 研究背景と目的")
    add_para(
        doc,
        "大脳白質病変は高血圧などとの関連が知られる脳の虚血性変化であり、脳卒中等のリスク把握に関係する。"
        "本研究では、健診データではなく頭部MRI画像そのものから病変グレードを判別し、撮像法・スライス位置・"
        "モデル構造が性能へ与える影響を明らかにすることを目指した。",
    )
    add_bullets(doc, [
        "MRI画像からgrade 0～4を推定するモデルを構築する。",
        "FLAIR（FL）、T1、T2など撮像法ごとの有効性を比較する。",
        "CNNとViTを公平な条件で比較し、ViT/Transformerが有効となる条件を明らかにする。",
        "実際のラベル単位である患者単位で、リーケージのない評価方法を確立する。",
        "少数の重症例を含むクラス不均衡下で、Accuracyだけでなくmacro-F1、AUC、少数クラスRecallを評価する。",
    ])

    add_heading(doc, "3. データとラベル構造")
    add_heading(doc, "3.1 MRIデータ", 2)
    add_bullets(doc, [
        "対象患者数：1,154人。",
        "撮像法：主にFLAIR（FL）とT1を使用。T2を含むモダリティ別設定も作成済み。",
        "患者1人あたり：FL/T1合計で概ね44～50枚のaxial画像。",
        "画像ファイルは患者ID・撮像法・撮像情報・axial番号を含む。データ作成時に重複回避用のランダム接尾辞が付加される。",
        "患者内の全スライスは同一ラベルであり、1,154患者すべてで患者内ラベル純度1.0を確認した。",
    ])
    add_heading(doc, "3.2 元の患者単位クラス分布", 2)
    add_table(doc, ["元grade", "患者数", "全体比", "備考"], [
        ["grade0", 557, "48.3%", "最多クラス"],
        ["grade1", 444, "38.5%", "中間クラス"],
        ["grade2", 106, "9.2%", "中等度"],
        ["grade3", 39, "3.4%", "重症"],
        ["grade4", 8, "0.7%", "極端な少数クラス"],
    ])
    add_note(
        doc,
        "独立標本数に関する注意",
        "grade4は画像枚数では数百枚に見えても、独立した患者は8人しかいない。同一患者の46枚は強く相関しており、"
        "46人分の情報ではない。このためgrade4を独立クラスとして学習・評価することは統計的に非常に不安定である。",
    )

    add_heading(doc, "4. 初期CNN研究と高精度結果")
    add_para(
        doc,
        "初期研究ではCNNを用い、image_size、axial範囲、撮像法の組合せを変えながら5クラス分類を行った。"
        "80×80px、axial 9～15、FL+T1条件で見かけ上のtest Accuracy約0.92を得た。元資料には各gradeのAUCとして"
        "grade0=0.9814、grade1=0.9800、grade2=0.9905、grade3=0.9977、grade4=0.9998が記録されている。",
    )
    add_note(
        doc,
        "後から判明した重大な問題",
        "この高精度は画像単位splitで得られ、同一患者の別スライスがtrainとvalidation/testへ分散していた。"
        "旧分割ではvalidation患者1,148人中1,146人がtrainにも存在しており、ほぼ100%の患者リーケージだった。"
        "したがって、この0.92は未見患者への汎化性能ではなく、患者固有特徴の記憶を強く含む参考値として扱う。",
        fill="FCE4D6",
    )

    add_heading(doc, "5. 患者リーケージの発見と再評価")
    add_para(
        doc,
        "ラベルが患者単位であることを確認したため、評価単位も患者へ合わせる必要があると判断した。"
        "データ準備処理にpatient splitを追加し、同一患者の全スライスを必ず同一splitへ配置した。",
    )
    add_table(doc, ["split", "画像数", "患者数", "患者重複"], [
        ["train", "39,883", "865", "0"],
        ["validation", "13,311", "289", "0"],
    ])
    add_bullets(doc, [
        "教師データにはFL/T1のpreprocess済みCSVを使用した。",
        "患者IDをファイル名先頭から抽出し、患者単位で層化分割した。",
        "validationの多数派grade0を常に予測した場合のAccuracyは約0.48である。",
        "以後、患者単位splitを研究の標準条件とした。",
    ])

    add_heading(doc, "6. 患者単位split・5クラス実験")
    add_para(
        doc,
        "患者リーケージを除去すると、5クラス分類は大きく難化した。特にgrade2～4はgrade0/1へ予測される傾向が強く、"
        "重症少数クラスが崩壊した。",
    )
    add_table(doc, ["モデル", "評価単位", "Accuracy", "macro-F1", "macro-AUC"], [
        ["CNN ResNet18", "slice", "0.495", "0.279", "0.654"],
        ["CNN ResNet18", "patient mean", "0.554", "0.258", "0.754"],
        ["ViT base", "slice", "0.444", "0.237", "0.563"],
        ["ViT base", "patient mean", "0.478", "0.204", "0.694"],
    ])
    add_para(
        doc,
        "患者平均集約によりAUCは改善したが、grade4は8患者しかなく、5クラス分類として安定した学習を行うには症例数が不足していた。",
    )

    add_heading(doc, "7. 3クラス化")
    add_para(
        doc,
        "病変なし・軽度・中等度以上という実用的区分と、クラス不均衡の緩和を目的に、grade2・3・4をgrade2+へ統合した。"
        "データは5クラス患者splitからhardlinkで作成し、患者splitを維持した。",
    )
    add_table(doc, ["クラス", "構成", "患者数（全体）"], [
        ["grade0", "元grade0", 557],
        ["grade1", "元grade1", 444],
        ["grade2+", "元grade2・3・4", 153],
    ])
    add_table(doc, ["モデル", "slice Acc", "slice F1", "slice AUC", "patient Acc", "patient F1", "patient AUC"], [
        ["CNN ResNet18", "0.521", "0.443", "0.634", "0.578", "0.489", "0.721"],
        ["ViT base", "0.452", "0.385", "0.563", "0.519", "0.411", "0.640"],
    ], font_size=8.2)
    add_para(
        doc,
        "3クラス化により、CNNの患者macro-F1は5クラス時の約0.26から0.489へ大幅に改善した。"
        "一方、旧ViT設定はCNNを下回った。",
    )

    add_heading(doc, "8. クラス不均衡対策")
    add_heading(doc, "8.1 過補正した過去実験", 2)
    add_para(
        doc,
        "weighted samplerと手動class weight（最大10倍）、さらにbest_metric=F1を同時に用いた実験では、"
        "少数クラスを過剰予測して全体性能が崩壊した。複数の補正法を重ねる危険性を確認した。",
    )
    add_heading(doc, "8.2 class-balanced loss単独", 2)
    add_table(doc, ["条件", "patient Acc", "macro-F1", "AUC", "grade2+ Recall"], [
        ["3class CNN baseline", "0.578", "0.489", "0.721", "0.179"],
        ["effective-beta重みのみ", "0.571", "0.475", "0.708", "0.154"],
    ])
    add_para(
        doc,
        "effective beta=0.9999の穏やかなclass-balanced loss単独でも改善せず、grade2+ Recallも低下した。"
        "3クラス化後の不均衡は約3.7:1であり、単純な重み付けは不要という否定結果を得た。",
    )

    add_heading(doc, "9. axial 9–15実験")
    add_para(
        doc,
        "病変が写りにくい周辺スライスを除外し、中央部のaxial 9～15へ限定した。CNNではAccuracyやAUCの大幅な上昇はなかったが、"
        "grade2+ Recallが大きく改善した。",
    )
    add_table(doc, ["CNN条件", "患者集約", "Acc", "macro-F1", "AUC", "grade2+ Recall"], [
        ["全axial", "mean", "0.578", "0.489", "0.721", "0.179"],
        ["全axial", "max confidence", "0.571", "0.519", "0.652", "0.282"],
        ["axial 9–15", "mean", "0.564", "0.496", "0.701", "0.410"],
        ["axial 9–15", "max confidence", "0.550", "0.496", "0.667", "0.513"],
    ])
    add_para(
        doc,
        "周辺スライスが患者ラベルに対するノイズとなっており、中央部の情報を重視することが重症例検出へ有効と解釈した。",
    )

    add_heading(doc, "10. ViT向け条件への転換")
    add_para(
        doc,
        "旧ViT実験はimage_size=80、patch16のため5×5トークンしかなく、vit_baseは1,154患者規模に対して大きすぎた。"
        "さらにRMSprop、augmentationなしというCNN再現条件をそのまま適用しており、ViTの評価として不公平だった。"
        "そこでViTに適した条件へ転換した。",
    )
    add_table(doc, ["項目", "旧ViT", "新DeiT-small"], [
        ["backbone", "vit_base_patch16", "deit_small_patch16_224"],
        ["パラメータ数", "約86M", "約21.7M"],
        ["入力", "80px（5×5 token）", "224px（14×14 token）"],
        ["optimizer", "RMSprop", "AdamW"],
        ["learning rate", "1e-4", "3e-5"],
        ["weight decay", "0", "0.05"],
        ["augmentation", "なし", "回転・flip・輝度・contrast・gamma"],
        ["事前学習", "ImageNet", "ImageNet/DeiT"],
    ])

    add_heading(doc, "11. 正則化DeiT-small実験")
    add_heading(doc, "11.1 非正則化DeiT-small 224", 2)
    add_para(
        doc,
        "224px・AdamW・augmentationを導入したが、epoch 1でval性能が最大となり、train Accuracyは約0.99まで上昇した一方、"
        "val lossは0.91から2.79へ増大した。明確な過学習を確認した。",
    )
    add_heading(doc, "11.2 過学習対策", 2)
    add_bullets(doc, [
        "epochs：50から30へ短縮。",
        "weight decay：0.05から0.1へ強化。",
        "cosine learning-rate schedulerを追加。",
        "label smoothing=0.1を追加。",
        "dropout=0.1、stochastic depth（drop path）=0.1を追加。",
        "augmentationをrotation 15°、brightness/contrast 0.2、gamma 0.8～1.2へ強化。",
    ])
    add_table(doc, ["指標", "非正則化", "正則化", "解釈"], [
        ["train Acc最終", "約0.990", "0.985", "記憶は依然強い"],
        ["val loss最終", "2.79", "1.52", "暴走は抑制"],
        ["best slice Acc", "0.543", "0.542", "ほぼ同等"],
        ["best slice F1", "0.469", "0.458", "改善なし"],
        ["best slice AUC", "0.686", "0.684", "改善なし"],
    ])
    add_para(
        doc,
        "スライス単位のピーク指標は改善しなかったが、後述する患者集約では正則化版の信頼度分布が有効に働き、"
        "max/top-k集約が大幅に改善した。",
    )

    figure = MRI_ROOT / "outputs/repro_vit_all_axial_patient_split_3class_deit_small_224_reg/figures/vit_learning_curves.png"
    if figure.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(figure), width=Inches(6.25))
        add_caption(doc, "図1　正則化DeiT-small 224の学習曲線（trainとvalidationの乖離を確認）")

    add_heading(doc, "12. 患者集約")
    add_heading(doc, "12.1 必要性", 2)
    add_para(
        doc,
        "正解ラベルは患者単位だが、モデルは各2Dスライスの確率を出力する。患者集約は、同一患者の約44～50枚の予測をまとめ、"
        "患者1人につき最終予測を1つ生成する処理である。これは評価時の後処理であり、現在のモデル学習自体はスライス単位である。",
    )
    add_heading(doc, "12.2 集約式", 2)
    add_bullets(doc, [
        "mean：p_patient = (1/N) Σ p_i。全スライス確率の平均。安定するが重要スライスの信号を薄める。",
        "max confidence：i* = argmax_i max_c p_i,c とし、p_patient = p_i*。最も自信の高い1枚を採用。",
        "attention confidence：各スライスの最大softmax確率を重みにして加重平均。",
        "attention entropy：予測分布の低entropyなスライスを重くする。実際のViT内部attentionではなく、確率entropyによる重み付け。",
        "top-k confidence：信頼度上位k枚の確率を平均。単一スライス依存を緩和するため、新たに実装。",
    ])
    add_note(
        doc,
        "患者集約はインチキか",
        "患者リーケージがなく、集約規則を最終test評価前に固定し、CNNにも同じ規則を適用すれば正当である。"
        "ただし現在は複数のpoolingを同一validationで比較しており、選択バイアスがある。現在値は開発結果であり、最終test結果ではない。",
    )

    add_heading(doc, "13. top-k集約と選択画像分析")
    add_heading(doc, "13.1 正則化ViTの患者単位結果", 2)
    add_table(doc, ["集約", "Accuracy", "macro-F1", "macro-AUC", "grade2+ Recall"], [
        ["mean", "0.581", "0.467", "0.764", "0.103"],
        ["attention entropy", "0.595", "0.510", "0.774", "0.179"],
        ["max confidence", "0.637", "0.634", "0.811", "0.769"],
        ["top-3", "0.616", "0.606", "0.811", "0.615"],
        ["top-5", "0.630", "0.631", "0.816", "0.615"],
    ])
    add_para(
        doc,
        "top-5はmax confidenceに近いmacro-F1を維持しながら、AUCは0.816で最大だった。"
        "したがって、max confidenceの好成績が単一の偶然の高信頼スライスだけによる可能性は低くなった。",
    )
    add_heading(doc, "13.2 top-5で選択された画像", 2)
    add_table(doc, ["分析項目", "結果"], [
        ["撮像法", "FL 86.0%、T1 14.0%"],
        ["axial 9～15", "全選択画像の60.5%"],
        ["頻出axial", "14、13、15、16、12"],
        ["grade2+患者での撮像法", "FL 92.8%、T1 7.2%"],
    ])
    add_para(
        doc,
        "モデルが主に中央付近のFL画像を選択しており、白質病変が確認しやすいと考えられる位置・撮像法と整合する。"
        "患者CSVには選択パス、撮像法、axial番号、信頼度を保存するよう実装した。",
    )
    add_note(
        doc,
        "ファイル名解析で発見した注意点",
        "データ準備処理は重複回避のため末尾へ「-<random>」を付加する。例「_15-0039」の0039はaxial番号ではなく"
        "ランダム接尾辞であり、axial番号はその直前の15である。評価コードを修正し、選択画像の位置を正しく記録した。",
    )

    add_heading(doc, "14. 元グレード別サブグループ解析")
    add_para(
        doc,
        "3クラス実験ではgrade2・3・4をgrade2+へ統合している。元ラベルへ戻して、各重症度がgrade2+として検出された割合を確認した。",
    )
    add_table(doc, ["元grade", "val患者数", "mean", "top-5", "max confidence"], [
        ["grade2", 27, "2/27 = 0.074", "14/27 = 0.519", "19/27 = 0.704"],
        ["grade3", 10, "2/10 = 0.200", "8/10 = 0.800", "9/10 = 0.900"],
        ["grade4", 2, "0/2 = 0.000", "2/2 = 1.000", "2/2 = 1.000"],
    ])
    add_para(
        doc,
        "情報のあるスライスを選択すると重症度が高いほど検出率が上昇し、視覚的に重症例が分かりやすいという直感と一致した。"
        "meanでは病変が写らない周辺・T1スライスによって信号が薄まったと考えられる。",
    )
    add_note(
        doc,
        "grade4に関する解釈",
        "grade4はvalidationに2人しかいないため、2/2という結果から一般化性能100%とは主張できない。"
        "また3クラスモデルはgrade4をgrade2+として検出しただけで、grade2・3・4を相互に識別したわけではない。"
        "真のgrade4分類には症例追加、患者単位交差検証、またはordinal regressionが必要である。",
    )

    add_heading(doc, "15. 公平比較CNNと単一split比較")
    add_para(
        doc,
        "ViT固有の効果と患者集約の効果を分離するため、ResNet18を正則化DeiT-smallと同じ条件で30epoch学習した。"
        "入力224px、batch size 16、AdamW、learning rate 3e-5、weight decay 0.1、cosine scheduler、label smoothing 0.1、"
        "augmentation、患者split、best-loss checkpoint、患者top-5を一致させた。ViT固有のdropout/drop path以外は公平な条件である。",
    )
    add_table(doc, ["289患者単一validation", "CNN", "ViT", "ViT－CNN"], [
        ["Accuracy", "0.5813", "0.6298", "+0.0484"],
        ["macro-F1", "0.5175", "0.6314", "+0.1139"],
        ["macro-AUC", "0.7702", "0.8160", "+0.0458"],
        ["grade2+ Recall", "0.538", "0.615", "+0.077"],
    ])
    add_table(doc, ["paired bootstrap（10,000回）", "ViT－CNN", "95% CI"], [
        ["Accuracy", "+0.0484", "0.0035～0.0969"],
        ["macro-F1", "+0.1139", "0.0579～0.1731"],
        ["macro-AUC", "+0.0458", "0.0193～0.0730"],
        ["balanced accuracy", "+0.0714", "0.0228～0.1245"],
    ])
    add_note(
        doc,
        "単一split結果の位置付け",
        "Accuracyのexact McNemar検定はp=0.09795であり、同一validationを用いてモデル・poolingを探索した選択バイアスもある。"
        "この大きなViT差は探索的所見とし、主結論には用いない。主解析は次節の全1,154患者5-fold OOF評価である。",
        fill="FFF2CC",
    )

    add_heading(doc, "16. 全axial・患者単位5-fold CV")
    add_heading(doc, "16.1 事前固定した解析条件", 2)
    add_para(
        doc,
        "単一validationへの依存を低減するため、既存train/validationの患者集合を統合し、"
        "StratifiedKFold（5分割、shuffle=True、random_state=1）で患者単位交差検証を構築した。"
        "各患者はvalidationにちょうど1回だけ含まれ、各foldのtrain/validation患者重複は0である。",
    )
    add_table(doc, ["項目", "固定条件"], [
        ["対象", "1,154患者・53,194画像"],
        ["クラス", "grade0=557、grade1=444、grade2+=153"],
        ["入力", "FL+T1・全axial・224px"],
        ["モデル", "DeiT-small / ResNet18"],
        ["学習", "各fold 30epoch、計10モデル"],
        ["患者集約", "best-loss checkpoint・top-5"],
        ["統計", "クラス層化paired patient bootstrap 10,000回"],
    ])
    add_note(
        doc,
        "全axialを主条件に固定",
        "axial 9～15の手動制限は解剖学的位置・撮像条件による境界の曖昧さがあり、再現性を損なう可能性がある。"
        "そのため主解析は全axialとし、axial 9～15は過去の参考ablationに限定した。",
        fill="EAF2F8",
    )

    add_heading(doc, "16.2 全1,154患者OOF結果", 2)
    add_table(doc, ["指標", "CNN（95% CI）", "ViT（95% CI）", "ViT－CNN（95% CI）"], [
        ["Accuracy", "0.6282（0.6040～0.6525）", "0.6352（0.6092～0.6603）", "+0.0069（－0.0156～0.0295）"],
        ["macro-F1", "0.6001（0.5682～0.6309）", "0.6134（0.5817～0.6441）", "+0.0133（－0.0135～0.0399）"],
        ["macro-AUC", "0.7891（0.7664～0.8111）", "0.8100（0.7889～0.8302）", "+0.0209（0.0062～0.0357）"],
        ["balanced accuracy", "0.5995", "0.6132", "+0.0137（－0.0112～0.0386）"],
    ], font_size=7.8)
    add_table(doc, ["クラスRecall", "CNN", "ViT", "ViT－CNN（95% CI）"], [
        ["grade0", "0.8636", "0.8402", "－0.0233（－0.0503～0.0036）"],
        ["grade1", "0.3468", "0.3851", "+0.0383（－0.0045～0.0811）"],
        ["grade2+", "0.5882", "0.6144", "+0.0261（－0.0261～0.0784）"],
    ])
    add_para(
        doc,
        "Accuracyのexact McNemar検定はp=0.5979で、ViTのみ正解92人、CNNのみ正解84人だった。"
        "Accuracy、macro-F1、balanced accuracy、各クラスRecallの差はいずれも95%信頼区間が0をまたいだ。"
        "macro-AUC差だけが0をまたがず、ViTの小さいランキング性能優位を支持した。",
    )
    add_note(
        doc,
        "5-fold後に固定した主結論",
        "DeiT-smallとResNet18の分類性能は同程度である。ViTの大幅なAccuracy/F1優位は再現しなかったが、"
        "macro-AUCでは小さいが再現性のある優位を示した。単一splitの大差にはsplit依存とselection optimismが含まれていたと判断する。",
        fill="E2F0D9",
    )

    add_heading(doc, "17. OOF予測の詳細解析")
    add_para(
        doc,
        "macro-AUC差の由来、少数重症クラスでの実用的判別、確率較正、モデル間の補完性を調べるため、"
        "固定済み1,154患者OOF確率を追加学習や閾値調整なしで解析した。信頼区間はクラス構成を維持したpaired bootstrap 10,000回で算出した。",
    )
    add_heading(doc, "17.1 クラス別ROC-AUCとPR-AUC", 2)
    add_table(doc, ["クラス", "指標", "CNN", "ViT", "ViT－CNN（95% CI）"], [
        ["grade0", "ROC-AUC", "0.7845", "0.8122", "+0.0276（0.0119～0.0437）"],
        ["grade0", "PR-AUC/AP", "0.7478", "0.7746", "+0.0268（0.0015～0.0528）"],
        ["grade1", "ROC-AUC", "0.6963", "0.7144", "+0.0181（－0.0037～0.0401）"],
        ["grade1", "PR-AUC/AP", "0.5671", "0.5803", "+0.0132（－0.0226～0.0477）"],
        ["grade2+", "ROC-AUC", "0.8866", "0.9035", "+0.0169（－0.0048～0.0400）"],
        ["grade2+", "PR-AUC/AP", "0.7097", "0.7104", "+0.0006（－0.0432～0.0471）"],
    ], font_size=7.8)
    add_para(
        doc,
        "全クラスでViTの点推定は同方向だったが、95%信頼区間が0をまたがなかったのはgrade0のROC-AUCとPR-AUCのみだった。"
        "したがってmacro-AUC優位の最も確実な由来はgrade0の順位付け改善であり、grade2+の検出優位性は確認されなかった。",
    )
    detail_figures = MRI_ROOT / "outputs/cv5_all_axial_3class/oof_detailed_analysis"
    add_figure(doc, detail_figures / "class_roc_curves.png", "図2　5-fold OOFのクラス別ROC曲線")
    add_figure(doc, detail_figures / "class_pr_curves.png", "図3　5-fold OOFのクラス別precision-recall曲線")

    add_heading(doc, "17.2 確率予測とキャリブレーション", 2)
    add_table(doc, ["指標（低いほど良い）", "CNN", "ViT", "ViT－CNN（95% CI）"], [
        ["Negative log-likelihood", "0.7894", "0.7532", "－0.0362（－0.0590～－0.0140）"],
        ["Multiclass Brier", "0.4864", "0.4627", "－0.0237（－0.0398～－0.0081）"],
        ["Top-label ECE", "0.0719", "0.0579", "－0.0140（－0.0345～0.0096）"],
    ])
    add_para(
        doc,
        "ViTはNLLとmulticlass Brier scoreで有意に良好であり、ランキングだけでなく確率予測全体の品質も改善した。"
        "ECEは数値的に良好だったが差は確定的でなかった。",
    )
    add_figure(doc, detail_figures / "top_label_reliability.png", "図4　5-fold OOFのtop-label reliability")
    add_figure(doc, detail_figures / "class_reliability.png", "図5　5-fold OOFのクラス別reliability")

    add_heading(doc, "17.3 混同行列とエラー重複", 2)
    add_table(doc, ["paired error category", "患者数", "全体比"], [
        ["両モデル正解", 641, "55.5%"],
        ["ViTのみ正解", 92, "8.0%"],
        ["CNNのみ正解", 84, "7.3%"],
        ["両モデル不正解", 337, "29.2%"],
    ])
    add_para(
        doc,
        "grade1患者444人のうち233人を両モデルが誤分類しており、grade0/1境界が主要なボトルネックだった。"
        "一方だけが正解した患者は176人存在し、ensembleや学習可能な患者集約の可能性を示すが、実際のensemble性能を保証するものではない。",
    )
    add_figure(doc, detail_figures / "confusion_matrices.png", "図6　5-fold OOFの患者単位混同行列")
    add_figure(doc, detail_figures / "error_overlap.png", "図7　CNNとViTの患者単位エラー重複", width=5.8)

    add_heading(doc, "17.4 元grade別検出とtop-5監査", 2)
    add_table(doc, ["元grade", "患者数", "CNN Recall", "ViT Recall", "ViT－CNN 95% CI"], [
        ["grade0", 557, "0.864", "0.840", "－0.050～0.004"],
        ["grade1", 444, "0.347", "0.385", "－0.005～0.081"],
        ["grade2", 106, "0.481", "0.500", "－0.047～0.094"],
        ["grade3", 39, "0.795", "0.846", "－0.051～0.154"],
        ["grade4", 8, "1.000", "1.000", "0.000～0.000"],
    ])
    add_para(
        doc,
        "元grade別のモデル差はいずれも確定的でなかった。grade4は両モデル8/8だが、症例数8人のため一般化性能100%を意味しない。"
        "top-5選択画像はCNNでFL 73.1%、ViTでFL 74.1%だった。axial 9～15の割合はCNN 41.7%、ViT 54.5%で、"
        "全axial入力を維持してもViTが中央付近をより多く選ぶ探索的傾向を認めた。これは医学的因果を示すものではない。",
    )

    add_heading(doc, "18. 実装した機能")
    add_table(doc, ["対象", "実装内容"], [
        ["データ準備", "patient split、3クラス化、患者5-fold hardlinkデータ、漏洩検証"],
        ["モデル", "ViT/DeiTのdrop_rate・drop_path_rate、ResNet18公平比較"],
        ["学習", "AdamW、cosine scheduler、label smoothing、指標別best checkpoint"],
        ["評価", "slice/patient評価、mean/max/attention/top-k、患者クラス確率保存"],
        ["監査", "選択スライスパス・撮像法・axial番号・信頼度、元grade復元"],
        ["CV自動化", "10学習の順次実行、完了判定、1,154患者OOF統合"],
        ["統計", "paired patient bootstrap、McNemar、class-wise ROC/PR、NLL/Brier/ECE"],
        ["可視化", "学習曲線、ROC/PR、reliability、混同行列、error overlap"],
        ["再現性", "YAML protocol、seed固定、全axial主解析の事前固定"],
    ], font_size=8.0)

    add_heading(doc, "19. 得られた知見")
    add_bullets(doc, [
        "患者単位ラベルのMRI研究では患者単位splitが必須であり、実効サンプル数は画像枚数ではなく患者数である。",
        "5クラスはgrade3・4の患者数が不足し、3クラス化が性能安定化に最も大きく寄与した。",
        "class weightは常に有効ではなく、複数の不均衡補正を重ねると過補正を起こす。",
        "手動axial 9～15は参考結果としてRecallを改善したが、主解析は曖昧さの少ない全axialへ固定した。",
        "ViTを公平に評価するには224px、小型DeiT、AdamW、augmentation、正則化が必要だった。",
        "患者top-5はmeanによる病変信号の希釈とmax confidenceの単一画像依存を緩和した。",
        "単一splitで見えたViTの大差は5-foldで再現せず、split依存とselection optimismを含んでいた。",
        "5-foldでは分類性能は同程度で、ViTの堅牢な優位はmacro-AUC、NLL、Brier scoreに限定された。",
        "ViTのAUC差は特にgrade0で明確で、grade2+ PR-AUCの優位は示されなかった。",
        "grade1が主要な誤分類源であり、ラベル境界・画質・症例多様性の監査が必要である。",
    ])

    add_heading(doc, "20. 失敗・否定結果")
    add_table(doc, ["試行", "結果", "学んだこと"], [
        ["画像単位split", "約0.92の高精度", "患者リーケージであり未見患者性能ではない"],
        ["5クラス患者split", "grade2～4が崩壊", "患者数不足と極端な不均衡"],
        ["旧ViT条件", "CNNを下回る", "80px・vit_base・RMSpropはViTに不適"],
        ["重み付けの重複", "少数クラス過剰予測", "補正法を重ねない"],
        ["class-balanced loss単独", "baselineより低下", "3クラス化後は不要"],
        ["正則化ViTのslice指標", "改善なし", "ラベル単位に患者評価を合わせる必要"],
        ["mean患者集約", "重症信号を希釈", "全画像を同じ重みで扱わない"],
        ["単一splitのViT大差", "5-foldで縮小", "探索結果を最終結論にしない"],
        ["grade2+ PR-AUC差", "ほぼ0", "重症検出のViT優位は未証明"],
    ], font_size=8.0)

    add_heading(doc, "21. 評価上の注意と限界")
    add_bullets(doc, [
        "5-fold OOFは単一split依存を抑えるが、モデルやtop-5は以前の289患者validationでの探索後に固定された。",
        "paired bootstrapは5個の学習済みfoldモデルを条件とした患者標本の不確実性であり、学習初期値やfold再分割の全変動を含まない。",
        "5-fold CVの割当seedは1種類であり、repeated CVは未実施である。",
        "外部施設・異なる撮像装置・異なる撮像期間による独立検証は未実施である。",
        "grade4は全体8人であり、8/8正解でも個別gradeの安定した性能を示さない。",
        "患者top-5は評価後処理であり、患者単位で直接学習するMIL/Transformerではない。",
        "現在のOOFを見ながら閾値・top-kを再調整すると再び選択バイアスが生じるため、主解析条件は変更しない。",
        "選択スライスはモデルの判断根拠候補であり、医学的因果説明には専門医による確認が必要である。",
    ])

    add_heading(doc, "22. 今後の計画")
    add_heading(doc, "22.1 最優先：外部検証", 2)
    add_bullets(doc, [
        "別施設または別撮像期間の患者で、現在のモデル・全axial・top-5を変更せず評価する。",
        "Accuracy/F1だけでなく、class-wise ROC-AUC・PR-AUC、NLL、Brier、ECEを報告する。",
        "撮像装置・画質・FL/T1構成の差による性能低下を確認する。",
    ])
    add_heading(doc, "22.2 外部データがない場合", 2)
    add_bullets(doc, [
        "異なる患者分割seedを2つ追加し、合計3反復のrepeated 5-fold CVを行う。",
        "同一30epoch protocolを変更せず、macro-AUC差とgrade0優位が反復間で再現するか確認する。",
        "grade1の両モデル誤分類患者を、予測を伏せて専門医が再監査し、境界症例・ラベル不一致・画質を評価する。",
    ])
    add_heading(doc, "22.3 モデル発展", 2)
    add_bullets(doc, [
        "手動axial範囲ではなく、全スライスembeddingから患者単位attention/MIL aggregatorを学習する。",
        "CNNとViTへ同じ患者aggregatorを適用し、encoder差を公平に評価する。",
        "slice位置embeddingとFL/T1 modality embeddingを追加する。",
        "新手法の選択には外側OOFを使用せず、nested CVのinner validationまたは新しい外部データを用いる。",
        "gradeの順序性を利用するordinal regressionを探索する。",
    ])

    add_heading(doc, "23. 結論")
    add_para(
        doc,
        "本研究では、画像単位splitによるほぼ100%の患者リーケージを発見し、患者単位評価へ研究設計を修正した。"
        "その後、3クラス化、ViT向け学習条件、正則化、患者top-k集約、公平なResNet18比較を段階的に実施し、"
        "最終的に全1,154患者を各1回だけvalidationとする全axial患者単位5-fold OOF評価を完了した。",
    )
    add_para(
        doc,
        "主解析では、CNNのAccuracy 0.6282・macro-F1 0.6001・macro-AUC 0.7891に対し、"
        "ViTはAccuracy 0.6352・macro-F1 0.6134・macro-AUC 0.8100だった。"
        "AccuracyとF1のモデル差は確定的でなく、macro-AUC差+0.0209の95%信頼区間0.0062～0.0357だけが0をまたがなかった。"
        "ViTはNLLとBrier scoreでも良好であり、確率予測品質の小さい優位を示した。",
    )
    add_para(
        doc,
        "したがって最終的な結論は、『全axial・患者単位5-fold CVにおいて、DeiT-smallはResNet18と同程度の分類性能を示し、"
        "macro-AUCと確率予測品質では小さいが再現性のある優位性を示した』である。"
        "ただし、その最も確実な由来はgrade0の識別であり、grade2+の検出優位性は確認されていない。"
        "この結論を一般化するには、固定済み条件による外部施設検証またはrepeated patient-level CVが必要である。",
    )

    doc.add_page_break()
    add_heading(doc, "付録A. 実験結果一覧")
    add_table(doc, ["段階", "モデル/条件", "患者集約", "Acc", "F1", "AUC", "grade2+ Recall"], [
        ["5class", "CNN ResNet18", "mean", "0.554", "0.258", "0.754", "―"],
        ["5class", "ViT base", "mean", "0.478", "0.204", "0.694", "―"],
        ["3class旧設定", "CNN全axial", "mean", "0.578", "0.489", "0.721", "0.179"],
        ["3class旧設定", "CNN全axial", "max", "0.571", "0.519", "0.652", "0.282"],
        ["3class旧設定", "ViT base", "mean", "0.519", "0.411", "0.640", "0.103"],
        ["3class旧設定", "ViT base", "attention entropy", "0.526", "0.426", "0.640", "0.128"],
        ["3class", "CNN weighted", "mean", "0.571", "0.475", "0.708", "0.154"],
        ["3class", "CNN axial 9–15", "mean", "0.564", "0.496", "0.701", "0.410"],
        ["3class", "CNN axial 9–15", "max", "0.550", "0.496", "0.667", "0.513"],
        ["3class 224", "DeiT-small非正則化", "mean", "0.581", "0.488", "0.760", "0.154"],
        ["3class 224", "DeiT-small非正則化", "max", "0.595", "0.571", "0.783", "0.846"],
        ["3class 224 reg", "DeiT-small", "mean", "0.581", "0.467", "0.764", "0.103"],
        ["3class 224 reg", "DeiT-small", "attention entropy", "0.595", "0.510", "0.774", "0.179"],
        ["3class 224 reg", "DeiT-small", "max", "0.637", "0.634", "0.811", "0.769"],
        ["3class 224 reg", "DeiT-small", "top-3", "0.616", "0.606", "0.811", "0.615"],
        ["3class 224 reg", "DeiT-small", "top-5", "0.630", "0.631", "0.816", "0.615"],
        ["公平単一split", "ResNet18", "top-5", "0.581", "0.518", "0.770", "0.538"],
        ["CV5 OOF主解析", "ResNet18", "top-5", "0.628", "0.600", "0.789", "0.588"],
        ["CV5 OOF主解析", "DeiT-small", "top-5", "0.635", "0.613", "0.810", "0.614"],
    ], font_size=7.6)

    add_heading(doc, "付録B. 主要ファイル")
    add_table(doc, ["区分", "workspace内パス", "内容"], [
        ["学習", "mri-vit-classification/src/train.py", "学習・指標別checkpoint・scheduler"],
        ["評価", "mri-vit-classification/src/evaluate.py", "患者集約・top-k・選択画像監査"],
        ["モデル", "mri-vit-classification/src/model.py", "ViT/ResNet18・dropout/drop path"],
        ["データ", "mri-vit-classification/src/dataset.py", "transform・augmentation・sampler"],
        ["データ準備", "mri-vit-classification/src/prepare_repro_vit_dataset.py", "患者split・axial絞込"],
        ["CVデータ準備", "mri-vit-classification/src/prepare_patient_cv_folds.py", "患者5-fold・hardlink・漏洩検証"],
        ["CV実行", "mri-vit-classification/scripts/run_patient_cv.py", "10学習と患者top-5評価"],
        ["CV統合", "mri-vit-classification/scripts/summarize_patient_cv.py", "1,154患者OOF統合"],
        ["統計比較", "mri-vit-classification/scripts/paired_patient_bootstrap.py", "paired bootstrap・McNemar"],
        ["OOF詳細解析", "mri-vit-classification/scripts/analyze_oof_predictions.py", "クラス別ROC/PR・較正・エラー"],
        ["正則化ViT config", "mri-vit-classification/config/config_repro_vit_all_axial_patient_split_3class_deit_small_224_reg.yaml", "現在のViT"],
        ["公平CNN config", "mri-vit-classification/config/config_repro_cnn_all_axial_patient_split_3class_resnet18_224_reg.yaml", "公平比較CNN"],
        ["CV固定protocol", "mri-vit-classification/config/generated_cv5_all_axial_3class/PROTOCOL.md", "全axial・top-5・30epoch"],
        ["旧結果", "mri-vit-classification/outputs/RESULTS_3class_patient_split_snapshot.md", "3class旧設定の凍結結果"],
        ["CV主結果", "mri-vit-classification/outputs/cv5_all_axial_3class/oof_summary/", "1,154患者OOFと主要bootstrap"],
        ["OOF詳細結果", "mri-vit-classification/outputs/cv5_all_axial_3class/oof_detailed_analysis/", "図表・患者エラーCSV・report"],
        ["top-5結果", "mri-vit-classification/outputs/repro_vit_all_axial_patient_split_3class_deit_small_224_reg/metrics/vit_eval_val_patient_top_k_confidence_k5_best_loss.json", "患者top-5指標"],
        ["top-5患者CSV", ".../metrics/vit_eval_val_patient_top_k_confidence_k5_best_loss_patients.csv", "選択画像メタデータ"],
        ["教師データ", "教師データ/labeled_image_list_FL_preprocess.csv・T1_preprocess.csv", "元grade・axial"],
    ], font_size=7.5)

    add_heading(doc, "付録C. 用語と数式")
    add_table(doc, ["用語", "意味"], [
        ["patient split", "同一患者の全画像を一つのsplitへまとめる分割"],
        ["patient leakage", "同一患者の別画像がtrainと評価側へ混在すること"],
        ["macro-F1", "各クラスF1を同じ重みで平均。少数クラス性能を反映"],
        ["macro-AUC", "多クラスone-vs-rest AUCのクラス平均"],
        ["mean pooling", "p_patient = (1/N) Σ p_i"],
        ["max confidence", "最もmax_c p_i,cが大きいスライスの予測を採用"],
        ["top-k confidence", "信頼度上位k枚の確率を平均"],
        ["MIL", "複数画像を一つの患者bagとして学習するMultiple Instance Learning"],
        ["ordinal regression", "gradeの順序関係0<1<2<3<4を利用する学習"],
    ])

    add_note(
        doc,
        "記録終了時点",
        "本書は2026年7月17日時点の研究状態を記録した。公平比較CNN、全axial患者単位5-fold CV、"
        "全1,154患者OOFのpaired bootstrapおよび詳細解析まで完了している。主解析条件と結論は本書へ固定し、"
        "今後は現在のOOFを見ながら再調整せず、外部検証またはrepeated patient-level CVへ進む。",
        fill="E2F0D9",
    )
    return doc


def main() -> None:
    doc = build_document()
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
