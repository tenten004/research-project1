# -*- coding: utf-8 -*-
"""X-ICT発表用の研究成果予稿をWordテンプレートから生成する。"""

from __future__ import annotations

import tempfile
from pathlib import Path
from typing import Iterable, Sequence

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


HERE = Path(__file__).resolve().parent
TEMPLATE_PATH = HERE / "XICT_MSWord_template.docx"
OUTPUT_PATH = HERE / "XICT予稿_T123001_浅川天夢_研究成果_20260804.docx"

FONT_JP_BODY = "Yu Mincho"
FONT_JP_TITLE = "Yu Gothic Medium"
FONT_EN = "Times New Roman"
COLOR_CNN = "#707070"
COLOR_VIT = "#2F6690"

JP_ABSTRACT = (
    "本研究では，脳MRI画像から大脳白質病変の重症度を予測するため，Vision Transformer（ViT）とCNNを"
    "患者単位で比較した．FLAIR・T1強調画像の全アキシャル53,194枚，1,154名をgrade 0，grade 1，"
    "grade 2以上の3クラスに統合し，患者重複のない5分割交差検証を実施した．各患者の高信頼5スライスを"
    "平均したout-of-fold評価では，ViTのmacro ROC-AUCは0.8100，CNNは0.7891で，差0.0209の95%信頼区間は"
    "0.0062–0.0357であった．一方，Accuracyとmacro-F1の差は有意でなかった．"
)

EN_ABSTRACT = (
    "We compare a Vision Transformer (ViT) with a convolutional neural network for patient-level "
    "classification of cerebral white matter lesion grades. A leakage-free five-fold cross-validation "
    "was performed on 53,194 FLAIR and T1-weighted axial images from 1,154 patients, using three classes: "
    "grade 0, grade 1, and grade 2 or higher. Patient probabilities were obtained by averaging the five "
    "most confident slices. The out-of-fold macro ROC-AUC was 0.8100 for ViT and 0.7891 for CNN; the paired "
    "difference was 0.0209 (95% CI: 0.0062–0.0357). Differences in accuracy and macro-F1 were inconclusive."
)


def set_run_font(run, size: float, *, bold: bool = False, title: bool = False) -> None:
    jp_font = FONT_JP_TITLE if title else FONT_JP_BODY
    run.font.name = FONT_EN
    run.font.size = Pt(size)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT_EN)
    rfonts.set(qn("w:hAnsi"), FONT_EN)
    rfonts.set(qn("w:eastAsia"), jp_font)


def clear_template_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def disable_font_embedding(doc: Document) -> None:
    settings = doc.settings.element
    element = settings.find(qn("w:embedTrueTypeFonts"))
    if element is None:
        element = OxmlElement("w:embedTrueTypeFonts")
        settings.append(element)
    element.set(qn("w:val"), "0")


def set_columns(section, count: int, space_twips: int = 425) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(count))
    cols.set(qn("w:space"), str(space_twips))
    if count == 1:
        cols.attrib.pop(qn("w:sep"), None)


def configure_section(section, columns: int) -> None:
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(25)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(8)
    set_columns(section, columns)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["normal"]
    normal.font.name = FONT_EN
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
    normal.element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_JP_BODY)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.space_after = Pt(0)


def add_centered(doc: Document, text: str, size: float, *, bold: bool = False,
                 title: bool = False, after: float = 0) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(after)
    run = paragraph.add_run(text)
    set_run_font(run, size, bold=bold, title=title)


def add_body(doc: Document, text: str, *, first_indent: bool = True,
             after: float = 2.5) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(after)
    if first_indent:
        paragraph.paragraph_format.first_line_indent = Pt(10.5)
    run = paragraph.add_run(text)
    set_run_font(run, 10.5)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(5 if level == 1 else 3)
    paragraph.paragraph_format.space_after = Pt(1.5)
    run = paragraph.add_run(text)
    set_run_font(run, 10.5, bold=True)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_borders(table, color: str = "666666", size: str = "4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]],
              widths_mm: Sequence[float], font_size: float = 7.4) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])

    for index, (header, width) in enumerate(zip(headers, widths_mm)):
        cell = table.rows[0].cells[index]
        cell.width = Mm(width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "D9E5F2")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        set_run_font(run, font_size, bold=True)

    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for column_index, (value, width) in enumerate(zip(values, widths_mm)):
            cell = cells[column_index]
            cell.width = Mm(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index % 2:
                set_cell_shading(cell, "F5F5F5")
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if column_index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            )
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            set_run_font(run, font_size)


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    set_run_font(run, 8.5)


def create_metric_figure(path: Path) -> None:
    labels = ["Accuracy", "Macro-F1", "Macro ROC-AUC", "Balanced Acc."]
    cnn = [0.6282, 0.6001, 0.7891, 0.5995]
    vit = [0.6352, 0.6134, 0.8100, 0.6132]
    positions = list(range(len(labels)))
    width = 0.34

    fig, axis = plt.subplots(figsize=(4.7, 2.45), dpi=220)
    axis.bar([p - width / 2 for p in positions], cnn, width, label="ResNet18", color=COLOR_CNN)
    axis.bar([p + width / 2 for p in positions], vit, width, label="ViT (DeiT-small)", color=COLOR_VIT)
    axis.set_ylim(0.5, 0.85)
    axis.set_ylabel("Score")
    axis.set_xticks(positions, labels, rotation=18, ha="right")
    axis.grid(axis="y", alpha=0.25, linewidth=0.6)
    axis.spines[["top", "right"]].set_visible(False)
    axis.legend(frameon=False, fontsize=7, ncol=2, loc="upper left")
    axis.tick_params(labelsize=7)
    axis.yaxis.label.set_size(8)
    for bars in axis.containers:
        axis.bar_label(bars, fmt="%.3f", padding=1, fontsize=6.5)
    fig.tight_layout(pad=0.6)
    fig.savefig(path, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def add_reference(doc: Document, index: int, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Pt(12)
    paragraph.paragraph_format.first_line_indent = Pt(-12)
    paragraph.paragraph_format.space_after = Pt(1)
    run = paragraph.add_run(f"[{index}] {text}")
    set_run_font(run, 8.3)


def validate_abstract_limits() -> None:
    jp_count = len(JP_ABSTRACT.replace(" ", ""))
    en_words = len(EN_ABSTRACT.split())
    if jp_count > 300:
        raise ValueError(f"Japanese abstract exceeds 300 characters: {jp_count}")
    if en_words > 100:
        raise ValueError(f"English abstract exceeds 100 words: {en_words}")
    print(f"Japanese abstract: {jp_count} characters")
    print(f"English abstract: {en_words} words")


def generate() -> Path:
    validate_abstract_limits()
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Template not found: {TEMPLATE_PATH}")

    doc = Document(TEMPLATE_PATH)
    clear_template_body(doc)
    disable_font_embedding(doc)
    configure_styles(doc)
    configure_section(doc.sections[0], columns=1)

    add_centered(
        doc,
        "Vision Transformerを用いた大脳白質病変グレード分類の患者単位評価",
        14,
        bold=True,
        title=True,
        after=3,
    )
    add_centered(doc, "浅川 天夢1)，石井 一夫1)", 12, after=1.5)
    add_centered(doc, "公立諏訪東京理科大学工学部情報応用工学科", 9, after=5)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(f"概要: {JP_ABSTRACT}")
    set_run_font(run, 9)

    add_centered(
        doc,
        "Patient-Level Evaluation of Cerebral White Matter Lesion Grade Classification Using Vision Transformer",
        13,
        bold=True,
        after=2,
    )
    add_centered(doc, "Ten Asakawa1), Kazuo Ishii1)", 10, after=1)
    add_centered(
        doc,
        "Department of Applied Information Engineering, Faculty of Engineering, Suwa University of Science",
        9,
        after=3,
    )
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(f"Abstract: {EN_ABSTRACT}")
    set_run_font(run, 9)

    body_section = doc.add_section(WD_SECTION.CONTINUOUS)
    configure_section(body_section, columns=2)

    add_heading(doc, "1. はじめに")
    add_body(
        doc,
        "大脳白質病変は脳MRIのT2強調画像やFLAIR画像で高信号を示し，加齢や高血圧と関連する．"
        "その進行は脳卒中や認知機能低下のリスクと関連するため，画像から重症度を客観的に評価する技術が求められる[1, 2]．",
    )
    add_body(
        doc,
        "先行研究ではCNNを用いたgrade 0からgrade 4の分類が行われた[1]．提案発表時には画像単位分割による"
        "Accuracy 0.9209を比較基準としていたが，患者IDを再点検すると，同一患者の別スライスが学習集合と評価集合に"
        "重複する構造が判明した．そこで本研究では評価単位を患者に改め，患者重複のない条件でVision Transformer"
        "（ViT）[3]とCNNを比較する．",
    )

    add_heading(doc, "2. 方法")
    add_heading(doc, "2.1 対象と分類課題", level=2)
    add_body(
        doc,
        "対象は脳ドック受診者1,154名で，FLAIR画像とT1強調画像の全アキシャル53,194枚を用いた．"
        "ラベルは患者単位で付与されている．grade 3および4は各39名，8名と少ないため，grade 2から4を統合し，"
        "grade 0（557名），grade 1（444名），grade 2以上（153名）の3クラス分類とした．",
    )
    add_table(
        doc,
        ["項目", "条件"],
        [
            ["対象", "1,154名，53,194画像"],
            ["入力", "FLAIR＋T1，全アキシャル"],
            ["クラス", "grade 0 / 1 / 2以上"],
            ["分割", "患者単位層化5-fold"],
            ["評価", "全患者のOOF予測"],
        ],
        [27, 52],
        font_size=8.1,
    )
    add_caption(doc, "表1: データと評価プロトコル")

    add_heading(doc, "2.2 比較条件と患者集約", level=2)
    add_body(
        doc,
        "ViTにはDeiT-small（patch 16，入力224×224）[4]，CNNにはResNet18を用いた．両者は同一fold，"
        "30 epoch，AdamW，学習率3×10−5，weight decay 0.1，cosine scheduling，label smoothing 0.1，"
        "同一データ拡張で学習した．各foldでvalidation lossが最小のモデルを採用した．",
    )
    add_body(
        doc,
        "患者ごとに全スライスを推論し，最大クラス確率が高い上位5スライスを選択してクラス確率を平均した．"
        "5-foldの各validation予測を連結し，全1,154名のout-of-fold（OOF）予測を得た．不確実性はクラス層化"
        "paired bootstrap 10,000回による95%信頼区間（CI）で評価した．",
    )

    add_heading(doc, "3. 結果")
    add_body(
        doc,
        "OOF評価を表2に示す．ViTは4指標すべてでCNNを上回った．ただし，95% CIが0を含まなかった差はmacro "
        "ROC-AUCのみであり，差は0.0209（95% CI: 0.0062–0.0357）であった．AccuracyのMcNemar検定は"
        "p=0.598であり，正解率の優位性は示されなかった．",
    )
    add_table(
        doc,
        ["指標", "CNN", "ViT", "ViT−CNN (95% CI)"],
        [
            ["Accuracy", "0.6282", "0.6352", "+0.0069 (−0.0156, 0.0295)"],
            ["Macro-F1", "0.6001", "0.6134", "+0.0133 (−0.0135, 0.0399)"],
            ["Macro AUC", "0.7891", "0.8100", "+0.0209 (0.0062, 0.0357)"],
            ["Balanced Acc.", "0.5995", "0.6132", "+0.0137 (−0.0112, 0.0386)"],
        ],
        [18, 13, 13, 35],
        font_size=6.8,
    )
    add_caption(doc, "表2: 患者単位5-fold OOF評価")

    with tempfile.TemporaryDirectory() as temp_dir:
        figure_path = Path(temp_dir) / "oof_metrics.png"
        create_metric_figure(figure_path)
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run().add_picture(str(figure_path), width=Mm(77))
    add_caption(doc, "図1: CNNとViTの患者単位OOF性能")

    add_body(
        doc,
        "クラス別再現率は，grade 0でCNN 0.8636，ViT 0.8402，grade 1で0.3468，0.3851，grade 2以上で"
        "0.5882，0.6144であった．いずれの差の95% CIも0を含んだ．特にgrade 1は両モデル共通の課題であり，"
        "境界症例の画像所見やラベルの揺らぎを臨床的に確認する必要がある．",
    )

    add_heading(doc, "4. 考察")
    add_body(
        doc,
        "患者単位の分割へ変更した結果，画像単位分割で得られていた高いAccuracyは再現せず，医療画像では"
        "患者リーケージを防ぐ評価設計が不可欠であることが確認された．一方，同一条件の5-fold比較でViTは"
        "macro ROC-AUCを改善した．Self-Attentionによる広域特徴の利用が，患者の重症度を順位づける能力に"
        "寄与した可能性がある．ただしAccuracyとmacro-F1の差は小さく，ViTが分類性能全般で優れるとは結論できない．",
    )
    add_body(
        doc,
        "本研究は全アキシャルを入力し，固定したtop-5集約で患者予測を得た．これにより手動の断面範囲指定を避けたが，"
        "スライス選択は確信度に基づく後処理であり，患者内の断面間関係を直接学習していない．また，本プロトコルは"
        "探索用validationでの検討後に固定したため，外部データに対する完全な未検証評価ではない．",
    )

    add_heading(doc, "5. おわりに")
    add_body(
        doc,
        "大脳白質病変の3クラス分類について，患者重複のない5-fold OOF評価でViTとCNNを比較した．ViTはmacro "
        "ROC-AUCで小さいが統計的に支持される改善を示した一方，Accuracyとmacro-F1の優位性は確認できなかった．"
        "今後は外部検証または反復交差検証を行うとともに，全スライスを患者単位で学習するMILやTransformer集約へ"
        "発展させ，grade 1の識別改善を検討する．",
    )

    add_heading(doc, "参考文献")
    add_reference(doc, 1, "竹村典晃，『畳み込みニューラルネットワークを用いた大脳白質病変のグレード予測』，公立諏訪東京理科大学卒業論文，2023．")
    add_reference(doc, 2, "S. Debette and H. S. Markus, “The clinical importance of white matter hyperintensities on brain magnetic resonance imaging,” BMJ, vol. 341, c3666, 2010.")
    add_reference(doc, 3, "A. Dosovitskiy et al., “An Image is Worth 16×16 Words: Transformers for Image Recognition at Scale,” ICLR, 2021.")
    add_reference(doc, 4, "H. Touvron et al., “Training data-efficient image transformers & distillation through attention,” ICML, pp. 10347–10357, 2021.")

    doc.core_properties.title = "Vision Transformerを用いた大脳白質病変グレード分類の患者単位評価"
    doc.core_properties.author = "浅川 天夢，石井 一夫"
    doc.core_properties.subject = "X-ICT発表予稿"
    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")
    return OUTPUT_PATH


if __name__ == "__main__":
    generate()